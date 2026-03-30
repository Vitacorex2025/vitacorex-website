from __future__ import annotations

import json
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document

from backend.core.config import settings
from backend.models.case_schema import CaseRecord, ContractReviewResult, FileType, StoredFile
from backend.services.contract_review_service import analyze_contracts
from backend.services.docx_extractor import extract_docx_text
from backend.services.file_classifier import classify_file
from backend.services.pdf_extractor import extract_pdf_text


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def contract_root() -> Path:
    root = settings.storage_path / "contract_intelligence"
    root.mkdir(parents=True, exist_ok=True)
    return root


def session_dir(session_id: str) -> Path:
    path = contract_root() / session_id
    path.mkdir(parents=True, exist_ok=True)
    return path


def session_meta_path(session_id: str) -> Path:
    return session_dir(session_id) / "session.json"


def _sanitize_filename(name: str) -> str:
    safe = "".join(ch for ch in str(name or "file") if ch.isalnum() or ch in {"-", "_", ".", " "}).strip()
    return safe or "file"


def _read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        payload = extract_pdf_text(path)
        return str(payload.get("text") or "")
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def _stored_file_payload(original_name: str, path: Path) -> dict[str, Any]:
    extracted_text = _read_text(path)
    file_type = classify_file(original_name, None, extracted_text)
    if file_type == "unknown":
        file_type = "contract_pdf"
    return {
        "id": f"contract_file_{uuid4().hex[:10]}",
        "original_name": original_name,
        "stored_name": path.name,
        "path": str(path),
        "size": path.stat().st_size,
        "extracted_text": extracted_text,
        "file_type": file_type,
    }


def _build_case_record(session: dict[str, Any]) -> CaseRecord:
    record = CaseRecord(
        id=f"contract_session_{session['session_id']}",
        workspace_id=session.get("workspace_id"),
        created_by_user_id=session.get("user_id"),
        reference=session["session_id"],
        title=session.get("title") or "Contract review session",
    )
    for item in session.get("contract_files", []):
        stored = StoredFile(
            id=item["id"],
            filename=item["stored_name"],
            original_filename=item["original_name"],
            workspace_id=session.get("workspace_id"),
            matter_id=record.id,
            content_type=None,
            file_type=FileType(str(item.get("file_type") or "contract_pdf")),
            storage_path=item["path"],
            size_bytes=int(item.get("size") or 0),
            extraction_version=0,
            extracted_text=item.get("extracted_text") or "",
        )
        stored.file_type = str(item.get("file_type") or "contract_pdf")
        record.source_files.append(stored)
    return record


def _session_summary(session: dict[str, Any]) -> str:
    contract_files = session.get("contract_files", [])
    if contract_files:
        return contract_files[0].get("original_name") or session.get("session_id")
    return session.get("session_id")


def create_session(contract_files: list[tuple[str, bytes]], supporting_files: list[tuple[str, bytes]] | None = None) -> dict[str, Any]:
    supporting_files = supporting_files or []
    try:
        from backend.services.auth_service import current_user_id, current_workspace_id

        active_user_id = current_user_id()
        active_workspace_id = current_workspace_id()
    except Exception:
        active_user_id = None
        active_workspace_id = None

    session_id = f"contract_{uuid4().hex[:12]}"
    directory = session_dir(session_id)
    contract_payloads: list[dict[str, Any]] = []
    support_payloads: list[dict[str, Any]] = []

    for index, (original_name, content) in enumerate(contract_files, start=1):
        stored = directory / f"contract_{index}_{_sanitize_filename(original_name)}"
        stored.write_bytes(content)
        contract_payloads.append(_stored_file_payload(original_name or stored.name, stored))

    for index, (original_name, content) in enumerate(supporting_files, start=1):
        stored = directory / f"support_{index}_{_sanitize_filename(original_name)}"
        stored.write_bytes(content)
        support_payloads.append(
            {
                "original_name": original_name or stored.name,
                "stored_name": stored.name,
                "path": str(stored),
                "size": stored.stat().st_size,
            }
        )

    session = {
        "session_id": session_id,
        "service_slug": "contract_intelligence",
        "created_at": _utcnow().isoformat(),
        "updated_at": _utcnow().isoformat(),
        "user_id": active_user_id,
        "workspace_id": active_workspace_id,
        "title": _session_summary({"contract_files": contract_payloads, "session_id": session_id}),
        "contract_files": contract_payloads,
        "supporting_files": support_payloads,
        "review": None,
        "status": "draft",
    }
    save_session(session)
    return session


def load_session(session_id: str) -> dict[str, Any]:
    path = session_meta_path(session_id)
    if not path.exists():
        raise FileNotFoundError(f"Contract session '{session_id}' does not exist.")
    session = json.loads(path.read_text(encoding="utf-8"))
    try:
        from backend.services.auth_service import current_user_id, current_workspace_id

        active_user_id = current_user_id()
        active_workspace_id = current_workspace_id()
        if active_user_id and session.get("user_id") and session.get("user_id") != active_user_id:
            raise FileNotFoundError(f"Contract session '{session_id}' does not exist.")
        if active_workspace_id and session.get("workspace_id") and session.get("workspace_id") != active_workspace_id:
            raise FileNotFoundError(f"Contract session '{session_id}' does not exist.")
    except FileNotFoundError:
        raise
    except Exception:
        pass
    return session


def save_session(session: dict[str, Any]) -> dict[str, Any]:
    session["updated_at"] = _utcnow().isoformat()
    session_meta_path(session["session_id"]).write_text(json.dumps(session, indent=2), encoding="utf-8")
    return session


def analyze_session(session_id: str) -> dict[str, Any]:
    session = load_session(session_id)
    record = _build_case_record(session)
    review = analyze_contracts(record)
    session["review"] = review.model_dump(mode="json")
    session["status"] = review.status or "completed"
    save_session(session)
    return session


def _hydrate_review(session: dict[str, Any]) -> ContractReviewResult:
    payload = session.get("review")
    if payload:
        return ContractReviewResult.model_validate(payload)
    analyzed = analyze_session(session["session_id"])
    return ContractReviewResult.model_validate(analyzed.get("review") or {})


def generate_draft(session_id: str) -> dict[str, Any]:
    session = load_session(session_id)
    review = _hydrate_review(session)
    if not review.findings:
        raise ValueError("No contract findings are available to support a revised draft.")

    directory = session_dir(session_id)
    revised_path = directory / "contract_review_revised.docx"
    memo_path = directory / "contract_review_change_memo.docx"

    doc = Document()
    doc.add_heading("Stronger Contract Draft", 0)
    doc.add_paragraph(
        "Human review required. This draft keeps the apparent business deal intact while replacing weak or missing clause mechanics with stronger fallback language."
    )
    doc.add_paragraph(f"Session: {session_id}")
    doc.add_heading("Recommended clause updates", level=1)
    for finding in review.findings:
        doc.add_heading(finding.title, level=2)
        doc.add_paragraph(f"Finding type: {finding.finding_kind} | Risk: {finding.risk_severity}")
        if finding.extracted_text:
            doc.add_paragraph(f"Current wording: {finding.extracted_text}")
        doc.add_paragraph(f"Why it matters: {finding.why_it_matters or finding.practical_implication or 'Review required.'}")
        doc.add_paragraph(f"Stronger language: {finding.fallback_language or finding.recommended_correction or 'Review required.'}")
        if finding.how_to_improve:
            doc.add_paragraph(f"Drafting note: {finding.how_to_improve}")
        if finding.jurisdiction_note:
            doc.add_paragraph(f"State-specific caution: {finding.jurisdiction_note}")
    doc.save(revised_path)

    memo = Document()
    memo.add_heading("Detailed Contract Change Memo", 0)
    memo.add_paragraph(
        "Human review notice: this memo explains what is strong, weak, missing, or dangerous, and why each stronger clause was recommended."
    )
    for finding in review.findings:
        memo.add_heading(finding.title, level=1)
        memo.add_paragraph(f"Finding type: {finding.finding_kind}")
        memo.add_paragraph(f"Risk severity: {finding.risk_severity}")
        memo.add_paragraph(f"Why it matters: {finding.why_it_matters or finding.practical_implication or 'Review required.'}")
        memo.add_paragraph(f"What can go wrong: {finding.what_can_go_wrong or 'Review required.'}")
        memo.add_paragraph(f"How to improve it: {finding.how_to_improve or finding.recommended_correction or 'Review required.'}")
        memo.add_paragraph(f"Legal / commercial rationale: {finding.legal_rationale or 'Review required.'}")
        memo.add_paragraph(f"Fallback wording: {finding.fallback_language or 'Review required.'}")
        if finding.jurisdiction_note:
            memo.add_paragraph(f"State-specific caution: {finding.jurisdiction_note}")
        if finding.extracted_text:
            memo.add_paragraph(f"Extracted text: {finding.extracted_text}")
    memo.save(memo_path)

    review.revised_contract_path = str(revised_path)
    review.revised_contract_download_url = f"/files/contract_intelligence/{session_id}/{revised_path.name}"
    review.change_memo_path = str(memo_path)
    review.change_memo_download_url = f"/files/contract_intelligence/{session_id}/{memo_path.name}"
    review.generated_at = _utcnow()

    session["review"] = review.model_dump(mode="json")
    session["status"] = "generated"
    save_session(session)
    return session


def delete_session(session_id: str) -> None:
    directory = contract_root() / session_id
    if not directory.exists():
        raise FileNotFoundError(f"Contract session '{session_id}' does not exist.")
    shutil.rmtree(directory, ignore_errors=False)


def list_sessions(limit: int = 20) -> list[dict[str, Any]]:
    sessions: list[dict[str, Any]] = []
    for path in contract_root().glob("*/session.json"):
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
            load_session(payload["session_id"])
            sessions.append(payload)
        except Exception:
            continue
    sessions.sort(key=lambda item: item.get("updated_at") or item.get("created_at") or "", reverse=True)
    return sessions[:limit]


def session_payload(session: dict[str, Any]) -> dict[str, Any]:
    review = session.get("review") or {}
    return {
        "session_id": session["session_id"],
        "service_slug": session.get("service_slug", "contract_intelligence"),
        "title": session.get("title") or _session_summary(session),
        "status": session.get("status", "draft"),
        "created_at": session.get("created_at"),
        "updated_at": session.get("updated_at"),
        "contract_files": session.get("contract_files", []),
        "supporting_files": session.get("supporting_files", []),
        "review": review,
        "generated_output": {
            "revised_contract_download_url": review.get("revised_contract_download_url"),
            "change_memo_download_url": review.get("change_memo_download_url"),
        },
    }
