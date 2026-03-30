
from __future__ import annotations

import re
from copy import deepcopy
from datetime import datetime
from decimal import Decimal
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from backend.models.case_schema import CaseRecord, GeneratedDocument
from backend.services.template_registry import resolve_template_path, strict_templates
from backend.services.validation_service import collect_document_blockers


def _money(value: Decimal | int | float | str | None) -> str:
    amount = Decimal(str(value or "0")).quantize(Decimal("0.01"))
    return f"${amount:,.2f}"


def _number_to_words_under_1000(n: int) -> str:
    units = [
        "Zero", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
        "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
        "Seventeen", "Eighteen", "Nineteen",
    ]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
    if n < 20:
        return units[n]
    if n < 100:
        return tens[n // 10] + ("" if n % 10 == 0 else f"-{units[n % 10]}")
    return units[n // 100] + " Hundred" + ("" if n % 100 == 0 else f" {_number_to_words_under_1000(n % 100)}")


def amount_to_words(value: Decimal | int | float | str | None) -> str:
    amount = Decimal(str(value or "0")).quantize(Decimal("0.01"))
    dollars = int(amount)
    cents = int((amount - Decimal(dollars)) * 100)
    parts: list[str] = []
    millions, remainder = divmod(dollars, 1_000_000)
    thousands, hundreds = divmod(remainder, 1_000)
    if millions:
        parts.append(_number_to_words_under_1000(millions) + " Million")
    if thousands:
        parts.append(_number_to_words_under_1000(thousands) + " Thousand")
    if hundreds or not parts:
        parts.append(_number_to_words_under_1000(hundreds))
    dollar_label = "Dollar" if dollars == 1 else "Dollars"
    cent_label = "Cent" if cents == 1 else "Cents"
    return f"({' '.join(parts)} {dollar_label} and {_number_to_words_under_1000(cents)} {cent_label})."

_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI"]


def _sequential_supported_count_labels(record: CaseRecord) -> list[str]:
    supported = [count for count in record.extracted_case_data.legal_counts if count.supported]
    results: list[str] = []
    for index, count in enumerate(supported, start=1):
        roman = _ROMAN[index - 1] if index - 1 < len(_ROMAN) else str(index)
        title = re.sub(r"^Count\s+[IVXLC]+\s+[–-]\s+", "", count.label, flags=re.IGNORECASE).strip()
        results.append(f"Count {roman} – {title}")
    return results



def _entity_descriptor(name: str | None, fallback: str | None = None) -> str:
    value = (name or "").upper()
    if " LLC" in f" {value}" or value.endswith("LLC"):
        return "a limited liability company"
    if " INC" in f" {value}" or value.endswith("INC") or " CORP" in f" {value}" or value.endswith("CORPORATION"):
        return "a corporation"
    if " LP" in f" {value}" or value.endswith("LP"):
        return "a limited partnership"
    return fallback or "a business entity"


def _flatten_address(address: str | None) -> str:
    cleaned = " ".join((address or "").replace("\n", " ").split()).strip()
    if cleaned:
        words = cleaned.split()
        for size in range(min(6, len(words) // 2), 1, -1):
            first = " ".join(words[:size]).lower()
            second = " ".join(words[size:size * 2]).lower()
            if first and first == second:
                words = words[size:]
                break
        if len(words) % 2 == 0:
            half = len(words) // 2
            if " ".join(words[:half]).lower() == " ".join(words[half:]).lower():
                words = words[:half]
        parts = words
        deduped: list[str] = []
        for token in parts:
            if len(deduped) >= 1 and deduped[-1].lower() == token.lower():
                continue
            deduped.append(token)
        cleaned = " ".join(deduped)
    return cleaned or "REVIEW REQUIRED — ADDRESS NOT EXTRACTED"


def _split_address_lines(address: str | None) -> tuple[str, str]:
    flat = _flatten_address(address)
    if "," in flat:
        first, second = flat.split(",", 1)
        return first.strip(), second.strip()
    parts = flat.split()
    if len(parts) > 6:
        midpoint = max(5, len(parts) // 2)
        return " ".join(parts[:midpoint]), " ".join(parts[midpoint:])
    return flat, ""


def _iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _replace_everywhere(doc: Document, old: str, new: str) -> None:
    if not old or old == new:
        return
    for paragraph in _iter_paragraphs(doc):
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)


def _replace_regex_everywhere(doc: Document, pattern: str, repl: str) -> None:
    compiled = re.compile(pattern, flags=re.IGNORECASE)
    for paragraph in _iter_paragraphs(doc):
        if compiled.search(paragraph.text):
            paragraph.text = compiled.sub(repl, paragraph.text)


def _replace_paragraph_prefix(doc: Document, prefix: str, new_text: str) -> bool:
    prefix_lower = prefix.lower()
    for paragraph in _iter_paragraphs(doc):
        if paragraph.text.strip().lower().startswith(prefix_lower):
            paragraph.text = new_text
            return True
    return False


def _paragraphs_list(doc: Document) -> list[Paragraph]:
    return list(doc.paragraphs)


def _insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    return new_para


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _delete_range(paragraphs: list[Paragraph], start: int, end_exclusive: int) -> None:
    for paragraph in paragraphs[start:end_exclusive]:
        _delete_paragraph(paragraph)


def _normalize_county(value: str | None) -> str | None:
    if not value:
        return None
    cleaned = " ".join(value.replace("County", "").split()).strip(" ,")
    return cleaned or None


def _extract_template_profile() -> dict[str, object]:
    profile: dict[str, object] = {
        "plaintiffs": set(),
        "defendants": set(),
        "guarantors": set(),
        "counties": set(),
        "addresses": set(),
        "amounts": set(),
    }

    for template in strict_templates():
        path = resolve_template_path(None, template.document_type)
        if not path or not Path(path).exists():
            continue
        doc = Document(str(path))
        texts = [paragraph.text.strip() for paragraph in _iter_paragraphs(doc) if paragraph.text.strip()]
        combined = "\n".join(texts)

        for amount in re.findall(r"\$\d[\d,]*\.\d{2}", combined):
            profile["amounts"].add(amount)

        for county in re.findall(r"IN AND FOR ([A-Z][A-Z ]+?) COUNTY, FLORIDA", combined):
            normalized = _normalize_county(county.title())
            if normalized:
                profile["counties"].add(normalized)

        for match in re.finditer(r"\b\d{1,6}[^,\n]{0,120},\s*[A-Za-z .'-]+,\s*(?:[A-Za-z ]+|[A-Z]{2})\s+\d{5}(?:-\d{4})?\b", combined):
            profile["addresses"].add(" ".join(match.group(0).split()))

        for index, text in enumerate(texts):
            if text == "Plaintiff(s)," and index > 0:
                profile["plaintiffs"].add(texts[index - 1].strip(" ,"))
            if text.startswith("Plaintiff:"):
                value = text.split(":", 1)[1].strip()
                if value:
                    profile["plaintiffs"].add(value)
            if text == "individually and as Personal Guarantor" and index > 0:
                guarantor = texts[index - 1].strip(" ,")
                if guarantor:
                    profile["guarantors"].add(guarantor)
            if text.startswith("To:"):
                value = text.split(":", 1)[1].strip()
                if value:
                    if "company" in template.document_type or template.document_type == "summons_company":
                        profile["defendants"].add(value)
                    elif template.document_type == "summons_individual":
                        profile["guarantors"].add(value)

        try:
            v_index = texts.index("v.")
        except ValueError:
            v_index = -1
        if v_index >= 0:
            for offset in range(v_index + 1, min(len(texts), v_index + 5)):
                line = texts[offset]
                if "Personal Guarantor" in line or line in {"Defendants.", "Defendant."}:
                    continue
                if ", a " in line.lower() or "limited liability company" in line.lower() or "corporation" in line.lower():
                    profile["defendants"].add(line.split(",", 1)[0].strip())
                    break

    return {
        "plaintiffs": sorted(profile["plaintiffs"]),
        "defendants": sorted(profile["defendants"]),
        "guarantors": sorted(profile["guarantors"]),
        "counties": sorted(profile["counties"]),
        "addresses": sorted(profile["addresses"]),
        "amounts": sorted(profile["amounts"]),
    }


@lru_cache(maxsize=1)
def template_profile() -> dict[str, object]:
    return _extract_template_profile()


def _case_number(record: CaseRecord) -> str:
    for note in record.notes:
        if note and str(note).lower().startswith("case no"):
            return str(note).split(":", 1)[-1].strip()
    return "__________________"


def _build_context(record: CaseRecord) -> dict[str, object]:
    data = record.extracted_case_data
    plaintiff_name = data.plaintiff.party_name or "REVIEW REQUIRED — PLAINTIFF"
    defendant_name = data.defendant_entity.party_name or "REVIEW REQUIRED — DEFENDANT ENTITY"
    guarantor_name = data.personal_guarantor.party_name if data.personal_guarantor and data.personal_guarantor.party_name else ""
    guarantor_present = bool(guarantor_name)

    venue_county = _normalize_county(data.venue.venue_county) or "REVIEW REQUIRED"
    amount = _money(data.amount_in_controversy)
    sunbiz = record.sunbiz_lookup
    verified_service_address = None
    if sunbiz and sunbiz.matched:
        verified_service_address = sunbiz.registered_agent_address or sunbiz.principal_address or sunbiz.mailing_address
    service_address = verified_service_address or data.service_address or data.business_address or data.defendant_entity.address or (data.personal_guarantor.address if data.personal_guarantor else None)
    company_address = verified_service_address or data.business_address or service_address
    guarantor_address = (data.personal_guarantor.address if data.personal_guarantor else None) or service_address

    company_line1, company_line2 = _split_address_lines(company_address)
    guarantor_line1, guarantor_line2 = _split_address_lines(guarantor_address)
    court_division = data.florida_court.court_division_label or "County Court"
    court_header = "IN THE CIRCUIT COURT" if "Circuit" in court_division else "IN THE COUNTY COURT"
    recommended_court = "Circuit Court" if "Circuit" in court_division else "County Court"
    entity_descriptor = _entity_descriptor(defendant_name, data.defendant_entity.entity_type)
    today = datetime.now()

    invoice_lines = [
        f"Invoice No. {item.invoice_number or 'REVIEW REQUIRED'} | Invoice Date: {item.invoice_date or 'REVIEW REQUIRED'} | Remaining Balance Due: {_money(item.balance_due)} | Source: {item.source_filename or item.source_sheet or item.source_kind or 'source'}"
        for item in data.invoice_summary.line_items
    ] or ["Invoice review required — no unique invoice-level balances were extracted."]

    selected_counts = _sequential_supported_count_labels(record)
    if not selected_counts:
        selected_counts = ["Count review required"]

    return {
        "plaintiff_name": plaintiff_name,
        "defendant_name": defendant_name,
        "defendant_caption": f"{defendant_name}, {entity_descriptor}" + (", and" if guarantor_present else ""),
        "guarantor_name": guarantor_name,
        "guarantor_caption_name": f"{guarantor_name}," if guarantor_present else "",
        "guarantor_descriptor": "individually and as Personal Guarantor" if guarantor_present else "",
        "guarantor_present": guarantor_present,
        "service_address_one_line": _flatten_address(service_address),
        "company_service_address_one_line": _flatten_address(company_address),
        "company_service_address_line1": company_line1,
        "company_service_address_line2": company_line2,
        "guarantor_service_address_one_line": _flatten_address(guarantor_address),
        "guarantor_service_address_line1": guarantor_line1,
        "guarantor_service_address_line2": guarantor_line2,
        "venue_county": venue_county,
        "venue_line": f"IN AND FOR {venue_county.upper()} COUNTY, FLORIDA",
        "court_header": court_header,
        "recommended_court": recommended_court,
        "case_number": _case_number(record),
        "amount": amount,
        "amount_words": amount_to_words(data.amount_in_controversy),
        "today_long": today.strftime("%B %d, %Y").replace(" 0", " "),
        "today_year": str(today.year),
        "governing_law": data.governing_law or "REVIEW REQUIRED — GOVERNING LAW NOT EXTRACTED",
        "payment_terms": data.payment_terms or "REVIEW REQUIRED — PAYMENT TERMS NOT EXTRACTED",
        "default_terms": data.default_terms or "REVIEW REQUIRED — DEFAULT TERMS NOT EXTRACTED",
        "interest_terms": data.interest_terms or "REVIEW REQUIRED — INTEREST TERMS NOT EXTRACTED",
        "venue_clause": data.venue.venue_clause or "REVIEW REQUIRED — VENUE CLAUSE NOT EXTRACTED",
        "venue_paragraph": data.venue.venue_paragraph or f"Venue is proper in {venue_county} County, Florida based on the strongest documented Florida nexus reflected in the uploaded record.",
        "invoice_lines": invoice_lines,
        "invoice_summary_sentence": f"The total unpaid principal balance currently due and owing is {amount}.",
        "legal_counts": selected_counts,
        "complaint_expected_count_headings": selected_counts,
        "proposed_judgment_line": (
            f"1. Plaintiff, {plaintiff_name}, shall recover from Defendants, {defendant_name} and {guarantor_name}, jointly and severally, the principal sum of {amount}."
            if guarantor_present else
            f"1. Plaintiff, {plaintiff_name}, shall recover from Defendant {defendant_name} the principal sum of {amount}."
        ),
        "affidavit_balance_sentence": (
            f"The total principal balance due and owing from Defendants is {amount}."
            if guarantor_present else
            f"The total principal balance due and owing from Defendant {defendant_name} is {amount}."
        ),
        "registered_agent_name": (sunbiz.registered_agent_name if sunbiz and sunbiz.matched else None) or "REVIEW REQUIRED — SUNBIZ VERIFICATION NEEDED",
        "registered_agent_address": _flatten_address((sunbiz.registered_agent_address if sunbiz and sunbiz.matched else None) or verified_service_address),
        "sunbiz_document_number": (sunbiz.florida_document_number if sunbiz and sunbiz.matched else None) or "REVIEW REQUIRED",
        "sunbiz_officer_line": ", ".join(filter(None, [((officer.officer_name or '') + (' — ' + officer.officer_title if officer.officer_title else '')).strip(' —') for officer in (sunbiz.officers if sunbiz and sunbiz.matched else [])])) or "REVIEW REQUIRED — SUNBIZ VERIFICATION NEEDED",
        "civil_cover_defendants": f"{defendant_name}; {guarantor_name}" if guarantor_present else defendant_name,
        "request_summons_sentence": "Plaintiff requests issuance of separate summons for each Defendant." if guarantor_present else "Plaintiff requests issuance of summons for the entity defendant only.",
        "wherefore_line": (
            f"WHEREFORE, Plaintiff respectfully requests entry of judgment against Defendants, jointly and severally, for the principal balance of {amount}, plus pre-judgment interest, court costs, post-judgment interest, and such other and further relief as this Court deems just and proper."
            if guarantor_present else
            f"WHEREFORE, Plaintiff respectfully requests entry of judgment against Defendant {defendant_name} for the principal balance of {amount}, plus pre-judgment interest, court costs, post-judgment interest, and such other and further relief as this Court deems just and proper."
        ),
        "damages_line": f"Plaintiff seeks recovery of the principal balance of {amount}, together with pre-judgment interest, court costs, post-judgment interest, and any other relief the Court deems proper.",
        "complaint_jurisdiction": f"The amount in controversy is {amount}. This action therefore falls within the jurisdiction of the {recommended_court} under Florida law.",
    }


def _apply_common_replacements(doc: Document, context: dict[str, object]) -> None:
    profile = template_profile()
    for value in profile["plaintiffs"]:
        _replace_everywhere(doc, str(value), str(context["plaintiff_name"]))
    for value in profile["defendants"]:
        _replace_everywhere(doc, str(value), str(context["defendant_name"]))
    for value in profile["guarantors"]:
        _replace_everywhere(doc, str(value), str(context["guarantor_name"]))
    for value in profile["amounts"]:
        _replace_everywhere(doc, str(value), str(context["amount"]))
    for value in profile["counties"]:
        _replace_everywhere(doc, f"{value} County", f"{context['venue_county']} County")
        _replace_everywhere(doc, f"COUNTY OF {value.upper()} COUNTY", f"COUNTY OF {str(context['venue_county']).upper()} COUNTY")
        _replace_everywhere(doc, f"IN AND FOR {str(value).upper()} COUNTY, FLORIDA", str(context["venue_line"]))

    for address in profile["addresses"]:
        _replace_everywhere(doc, str(address), str(context["service_address_one_line"]))

    _replace_regex_everywhere(doc, r"CASE NO:\s*_+", f"CASE NO: {context['case_number']}")
    _replace_regex_everywhere(doc, r"CASE NO:\s*[^\n]+", f"CASE NO: {context['case_number']}")
    _replace_paragraph_prefix(doc, "IN THE COUNTY COURT", str(context["court_header"]))
    _replace_paragraph_prefix(doc, "IN THE CIRCUIT COURT", str(context["court_header"]))
    _replace_paragraph_prefix(doc, "IN AND FOR", str(context["venue_line"]))


def _refresh_caption(doc: Document, context: dict[str, object]) -> None:
    paragraphs = _paragraphs_list(doc)
    plaintiff_marker = None
    v_index = None
    defendant_marker = None

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text == "Plaintiff(s),":
            plaintiff_marker = index
        elif text == "v.":
            v_index = index
        elif text in {"Defendants.", "Defendant."}:
            defendant_marker = index
            break

    if plaintiff_marker is not None and plaintiff_marker > 0:
        paragraphs[plaintiff_marker - 1].text = f"{context['plaintiff_name']},"

    if v_index is not None and defendant_marker is not None and defendant_marker > v_index:
        original_end = paragraphs[defendant_marker]
        original_end.text = "Defendants." if context["guarantor_present"] else "Defendant."
        _delete_range(paragraphs, v_index + 1, defendant_marker)
        anchor = paragraphs[v_index]
        anchor = _insert_after(anchor, str(context["defendant_caption"]))
        if context["guarantor_present"]:
            anchor = _insert_after(anchor, str(context["guarantor_name"]) + ",")
            _insert_after(anchor, str(context["guarantor_descriptor"]))

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not context["guarantor_present"] and (
            text == "individually and as Personal Guarantor"
            or text.endswith("Personal Guarantor")
            or text.lower().startswith("defendant:") and "guarantor" in text.lower()
        ):
            _delete_paragraph(paragraph)


def _delete_section_by_heading(doc: Document, heading_text: str, stop_prefixes: tuple[str, ...]) -> None:
    paragraphs = list(doc.paragraphs)
    in_section = False
    for paragraph in list(paragraphs):
        text = paragraph.text.strip()
        if not text and in_section:
            _delete_paragraph(paragraph)
            continue
        if text.upper().startswith(heading_text.upper()):
            in_section = True
        if in_section:
            if text != heading_text and any(text.upper().startswith(prefix.upper()) for prefix in stop_prefixes):
                in_section = False
                continue
            _delete_paragraph(paragraph)


def _apply_demand_letter(doc: Document, context: dict[str, object]) -> None:
    _replace_paragraph_prefix(doc, "Date:", f"Date: {context['today_long']}")
    _replace_paragraph_prefix(doc, "Re:", f"Re: Outstanding Balance Owed to {context['plaintiff_name']}")
    _replace_paragraph_prefix(doc, "Dear", f"Dear {context['guarantor_name'] or context['defendant_name']}:")
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Attn:"):
            paragraph.text = f"Attn: {context['guarantor_name'] or context['defendant_name']}"
        elif "total outstanding principal balance" in text.lower():
            paragraph.text = f"As of this date, the total outstanding principal balance is {context['amount']}."
    paragraphs = _paragraphs_list(doc)
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().startswith("Via Email"):
            if idx + 1 < len(paragraphs):
                paragraphs[idx + 1].text = str(context["defendant_name"])
            if idx + 2 < len(paragraphs) and paragraphs[idx + 2].text.strip().startswith("Attn:"):
                paragraphs[idx + 2].text = f"Attn: {context['guarantor_name'] or context['defendant_name']}"
            if idx + 3 < len(paragraphs):
                paragraphs[idx + 3].text = str(context["service_address_one_line"])
            break


def _apply_complaint(doc: Document, context: dict[str, object]) -> None:
    paragraphs = _paragraphs_list(doc)
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        lowered = text.lower()
        if re.match(r"^\d+\.\s+the amount in controversy", lowered):
            paragraph.text = re.sub(r"^\d+\.\s*", "5. ", str(context["complaint_jurisdiction"]))
        elif re.match(r"^\d+\.\s+venue", lowered):
            paragraph.text = re.sub(r"^\d+\.\s*", "6. ", str(context["venue_paragraph"]))
        elif "may be served at" in lowered:
            paragraph.text = re.sub(r"^\d+\.\s*", "4. ", f"Defendants may be served at {context['service_address_one_line']}.")
        elif lowered.startswith("wherefore"):
            paragraph.text = str(context["wherefore_line"])
        elif "plaintiff seeks recovery of the principal balance" in lowered:
            paragraph.text = str(context["damages_line"])

    expected_headings = list(context.get("complaint_expected_count_headings") or [])
    count_heading_indexes = [idx for idx, paragraph in enumerate(paragraphs) if re.match(r"^COUNT\s+[IVXLC]+", paragraph.text.strip(), flags=re.IGNORECASE)]
    for idx, paragraph_index in enumerate(count_heading_indexes):
        paragraph = paragraphs[paragraph_index]
        if idx < len(expected_headings):
            paragraph.text = expected_headings[idx].upper()
        else:
            next_breaks = count_heading_indexes[idx + 1:] + [len(paragraphs)]
            end_index = next_breaks[0]
            _delete_range(paragraphs, paragraph_index, end_index)
            break

    if not context["guarantor_present"]:
        _delete_section_by_heading(doc, "COUNT II", ("COUNT III", "COUNT IV", "DAMAGES", "WHEREFORE"))
        for paragraph in list(doc.paragraphs):
            lowered = paragraph.text.lower()
            if "personal guarantor" in lowered or ("guarantor" in lowered and paragraph.text.strip().startswith(("3.", "count ii"))):
                _delete_paragraph(paragraph)

    wherefore_seen = False
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().lower().startswith("wherefore"):
            if wherefore_seen:
                _delete_paragraph(paragraph)
            else:
                wherefore_seen = True
                paragraph.text = str(context["wherefore_line"])
    _refresh_caption(doc, context)


def _apply_civil_cover(doc: Document, context: dict[str, object]) -> None:
    _replace_paragraph_prefix(doc, "Plaintiff:", f"Plaintiff: {context['plaintiff_name']}")
    _replace_paragraph_prefix(doc, "Defendants:", f"Defendants: {context['civil_cover_defendants']}")
    _replace_paragraph_prefix(doc, "Amount in Controversy:", f"Amount in Controversy: {context['amount']}")
    _replace_paragraph_prefix(doc, "Recommended Court:", f"Recommended Court: {context['recommended_court']}")
    _replace_paragraph_prefix(doc, "Venue:", f"Venue: {context['venue_county']} County")
    if not context["guarantor_present"]:
        _replace_everywhere(doc, " / Guaranty", "")
        _replace_everywhere(doc, "/ Guaranty", "")
    _refresh_caption(doc, context)


def _apply_request_for_summons(doc: Document, context: dict[str, object]) -> None:
    paragraphs = _paragraphs_list(doc)
    service_heading_indices = [idx for idx, p in enumerate(paragraphs) if p.text.strip().lower().startswith("service address")]
    defendant_line_count = 0
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text.startswith("1. Defendant:"):
            paragraph.text = f"1. Defendant: {context['defendant_name']}"
            defendant_line_count += 1
        elif text.startswith("2. Defendant:"):
            if context["guarantor_present"]:
                paragraph.text = f"2. Defendant: {context['guarantor_name']}"
                defendant_line_count += 1
            else:
                _delete_paragraph(paragraph)
        elif "separate summons" in text.lower():
            paragraph.text = str(context["request_summons_sentence"])

    for idx, marker in enumerate(service_heading_indices):
        # first service address belongs to company, second to guarantor
        target_line1 = context["company_service_address_line1"] if idx == 0 else context["guarantor_service_address_line1"]
        target_line2 = context["company_service_address_line2"] if idx == 0 else context["guarantor_service_address_line2"]
        if marker + 1 < len(paragraphs):
            paragraphs[marker + 1].text = str(target_line1)
        if marker + 2 < len(paragraphs):
            if idx == 1 and not context["guarantor_present"]:
                _delete_paragraph(paragraphs[marker + 2])
            else:
                paragraphs[marker + 2].text = str(target_line2 or "")
    _refresh_caption(doc, context)


def _apply_summons(doc: Document, context: dict[str, object], individual: bool = False) -> None:
    addressee = context["guarantor_name"] if individual else context["defendant_name"]
    address_line = context["guarantor_service_address_one_line"] if individual else context["company_service_address_one_line"]
    _replace_paragraph_prefix(doc, "To:", f"To: {addressee}")
    paragraphs = _paragraphs_list(doc)
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().startswith("To:"):
            if idx + 1 < len(paragraphs):
                paragraphs[idx + 1].text = str(address_line)
            break
    _refresh_caption(doc, context)


def _apply_motion_for_default(doc: Document, context: dict[str, object]) -> None:
    _refresh_caption(doc, context)


def _apply_proposed_final_judgment(doc: Document, context: dict[str, object]) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("1. Plaintiff,"):
            paragraph.text = str(context["proposed_judgment_line"])
            break
    _refresh_caption(doc, context)


def _apply_debt_calculation_sheet(doc: Document, context: dict[str, object]) -> None:
    for paragraph in doc.paragraphs:
        lowered = paragraph.text.lower()
        if "total principal balance" in lowered or "currently due and owing" in lowered:
            paragraph.text = str(context["invoice_summary_sentence"])
        elif "defendants:" in lowered:
            paragraph.text = (
                f"Defendants: {context['defendant_name']} and {context['guarantor_name']}"
                if context["guarantor_present"]
                else f"Defendant: {context['defendant_name']}"
            )

    paragraphs = _paragraphs_list(doc)
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().startswith("Invoice Schedule"):
            cursor = paragraph
            while True:
                next_p = cursor._p.getnext()
                if next_p is None:
                    break
                next_para = Paragraph(next_p, cursor._parent)
                next_text = next_para.text.strip()
                if not next_text or next_text.startswith("Total") or next_text.startswith("Amount") or next_text.startswith("Further Affiant") or next_text.startswith("Sworn"):
                    break
                if "invoice" in next_text.lower() or "$" in next_text or "|" in next_text:
                    _delete_paragraph(next_para)
                else:
                    break
            for line in context["invoice_lines"]:
                cursor = _insert_after(cursor, line)
            break
    _refresh_caption(doc, context)


def _apply_affidavit(doc: Document, context: dict[str, object]) -> None:
    for paragraph in doc.paragraphs:
        lowered = paragraph.text.lower()
        if lowered.startswith("county of "):
            paragraph.text = f"COUNTY OF {str(context['venue_county']).upper()} COUNTY"
        elif "authorized representative of" in lowered:
            paragraph.text = re.sub(r"of\s+.+?,", f"of {context['plaintiff_name']},", paragraph.text)
        elif "total principal balance due and owing" in lowered:
            paragraph.text = str(context["affidavit_balance_sentence"])
        elif "invoice no." in lowered and "balance due" in lowered:
            paragraph.text = context["invoice_lines"][0]
    _refresh_caption(doc, context)


def _apply_exhibit_index(doc: Document, context: dict[str, object]) -> None:
    if not context["guarantor_present"]:
        for paragraph in list(doc.paragraphs):
            if "Exhibit C" in paragraph.text and "Guarant" in paragraph.text:
                _delete_paragraph(paragraph)
    _refresh_caption(doc, context)


def _apply_doc_specific_changes(doc: Document, document_type: str, context: dict[str, object]) -> None:
    if document_type == "demand_letter":
        _apply_demand_letter(doc, context)
    elif document_type == "complaint":
        _apply_complaint(doc, context)
    elif document_type == "civil_cover_sheet":
        _apply_civil_cover(doc, context)
    elif document_type == "request_for_issuance_of_summons":
        _apply_request_for_summons(doc, context)
    elif document_type == "summons_company":
        _apply_summons(doc, context, individual=False)
    elif document_type == "summons_individual":
        _apply_summons(doc, context, individual=True)
    elif document_type == "motion_for_clerks_default":
        _apply_motion_for_default(doc, context)
    elif document_type == "proposed_final_judgment":
        _apply_proposed_final_judgment(doc, context)
    elif document_type == "debt_calculation_sheet":
        _apply_debt_calculation_sheet(doc, context)
    elif document_type == "affidavit_of_amount_due":
        _apply_affidavit(doc, context)
    elif document_type == "exhibit_index":
        _apply_exhibit_index(doc, context)
    else:
        _refresh_caption(doc, context)


def _generated_doc_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for paragraph in _iter_paragraphs(doc):
        value = paragraph.text.strip()
        if value:
            parts.append(value)
    return "\n".join(parts)


def _qa_generated_document(record: CaseRecord, document_type: str, path: Path) -> list[str]:
    issues: list[str] = []
    text = _generated_doc_text(path)
    data = record.extracted_case_data
    amount = _money(data.amount_reconciliation.approved_total or data.amount_in_controversy)
    if "REVIEW REQUIRED" in text:
        issues.append("Document still contains REVIEW REQUIRED placeholders.")
    if document_type == "complaint":
        if text.count("WHEREFORE") != 1:
            issues.append("Complaint contains duplicated or missing WHEREFORE paragraph.")
        if data.defendant_entity.party_name and data.defendant_entity.party_name not in text:
            issues.append("Complaint is missing the debtor entity name.")
        if data.personal_guarantor and data.personal_guarantor.party_name and data.personal_guarantor.party_name not in text:
            issues.append("Complaint is missing the guarantor name.")
        if (
            re.search(r"Defendant\s{2,}(?:personally|exercised|is\b)", text, flags=re.IGNORECASE)
            or re.search(r"Defendant\s*(?:,|\.)\s*(?:personally|exercised|is\b|$)", text, flags=re.IGNORECASE)
        ):
            issues.append("Complaint contains a blank defendant or guarantor placeholder.")
        if amount not in text:
            issues.append("Complaint amount does not match the reconciled approved amount.")
    return issues


def generate_packet_documents(
    record: CaseRecord,
    output_dir: str | Path,
    selected_document_types: list[str] | None = None,
    progress_callback=None,
) -> list[GeneratedDocument]:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    context = _build_context(record)
    selected = set(selected_document_types or [template.document_type for template in strict_templates()])

    generated: list[GeneratedDocument] = []
    total = len(selected)
    completed = 0

    for template in strict_templates():
        if template.document_type not in selected:
            continue

        if progress_callback:
            progress_callback(template.document_type, template.output_filename, "generating", completed, total, None)

        if template.document_type == "summons_individual" and not context["guarantor_present"]:
            doc = GeneratedDocument(
                document_name=template.output_filename,
                document_path=None,
                document_type=template.document_type,
                download_url=None,
                status="not_applicable",
                blocking_reason="No individual guarantor was extracted from the uploaded record.",
            )
            generated.append(doc)
            completed += 1
            if progress_callback:
                progress_callback(template.document_type, template.output_filename, doc.status, completed, total, doc)
            continue

        document_blockers = collect_document_blockers(record, template.document_type)
        if document_blockers:
            doc = GeneratedDocument(
                document_name=template.output_filename,
                document_path=None,
                document_type=template.document_type,
                download_url=None,
                status="blocked",
                blocking_reason=" • ".join(document_blockers),
            )
            generated.append(doc)
            completed += 1
            if progress_callback:
                progress_callback(template.document_type, template.output_filename, doc.status, completed, total, doc)
            continue

        template_path = resolve_template_path(record, template.document_type)
        if template_path is None or not Path(template_path).exists():
            doc = GeneratedDocument(
                document_name=template.output_filename,
                document_path=None,
                document_type=template.document_type,
                download_url=None,
                status="template_missing",
                blocking_reason=f"Missing controlling template: {template.output_filename}",
            )
            generated.append(doc)
            completed += 1
            if progress_callback:
                progress_callback(template.document_type, template.output_filename, doc.status, completed, total, doc)
            continue

        docx = Document(str(template_path))
        _apply_common_replacements(docx, context)
        _apply_doc_specific_changes(docx, template.document_type, context)

        output_path = output_dir / template.output_filename
        docx.save(output_path)
        qa_issues = _qa_generated_document(record, template.document_type, output_path)
        if qa_issues:
            try:
                output_path.unlink()
            except Exception:
                pass
            doc = GeneratedDocument(
                document_name=template.output_filename,
                document_path=None,
                document_type=template.document_type,
                download_url=None,
                status="failed_qa",
                blocking_reason=" • ".join(qa_issues),
            )
        else:
            doc = GeneratedDocument(
                document_name=output_path.name,
                document_path=str(output_path),
                document_type=template.document_type,
                download_url=f"/files/cases/{record.id}/output/{output_path.name}",
                status="generated",
            )
        generated.append(doc)
        completed += 1
        if progress_callback:
            progress_callback(template.document_type, template.output_filename, doc.status, completed, total, doc)
    return generated
