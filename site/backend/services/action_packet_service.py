from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from backend.models.case_schema import ActionPacketOutput, CaseRecord
from backend.services.case_state import active_source_files
from backend.services.case_store import ensure_case_dirs
from backend.services.download_ledger import visible_entries
from backend.services.irs_notice_parser import parse_irs_notice
from backend.services.invoice_recovery_parser import parse_invoice_recovery_matter
from backend.services.recovery_routing_engine import determine_recovery_route
from backend.services.recovery_workspace_service import build_reconciliation_center, refresh_workspace_modules
from backend.services.validation_service import assess_packet_readiness


def _existing_outputs(record: CaseRecord) -> list[ActionPacketOutput]:
    outputs: list[ActionPacketOutput] = []
    for entry in visible_entries(record):
        outputs.append(
            ActionPacketOutput(
                type=entry.category or "generated_output",
                label=entry.file_name,
                status="generated",
                download_url=entry.download_url,
                file_path=entry.file_path,
                owner_kind=entry.owner_kind,
                owner_id=entry.owner_id,
                service_slug=entry.service_slug,
            )
        )
    return outputs


def _money(value: object | None) -> str:
    try:
        number = float(value or 0)
    except Exception:
        return "-"
    return f"${number:,.2f}"


def _case_file_url(record: CaseRecord, path: Path) -> str:
    return f"/files/cases/{record.id}/output/{path.name}"


def _fallback_invoice_payload(record: CaseRecord, payload: dict[str, object]) -> dict[str, object]:
    needs_repair = str(payload.get("supported_balance") or "0.00") == "0.00" or not payload.get("invoice_count")
    if not needs_repair:
        return payload

    texts: list[str] = []
    for item in active_source_files(record):
        text = str(item.extracted_text or "").strip()
        if text:
            texts.append(text)
            continue
        storage_path = Path(str(item.storage_path or ""))
        if storage_path.exists() and storage_path.suffix.lower() in {".txt", ".md"}:
            texts.append(storage_path.read_text(encoding="utf-8", errors="ignore"))
        else:
            texts.append(f"{item.original_filename} {item.file_type}")

    if not texts:
        return payload

    repaired = parse_invoice_recovery_matter(texts)
    merged = dict(payload)
    for key, value in repaired.items():
        if value not in (None, "", [], {}):
            merged[key] = value
    workflow_payload = dict(record.workflow_payload or {})
    workflow_payload["invoice_recovery"] = merged
    record.workflow_payload = workflow_payload
    return merged


def _write_simple_pdf(path: Path, title: str, lines: list[str]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 54
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(48, y, title[:90])
    y -= 24
    pdf.setFont("Helvetica", 10)
    for line in lines:
        if y < 56:
            pdf.showPage()
            y = height - 54
            pdf.setFont("Helvetica", 10)
        pdf.drawString(48, y, str(line)[:110])
        y -= 15
    pdf.save()


def _invoice_primary_output_spec(route_status: str) -> tuple[str, str, str, str]:
    if route_status == "ready_for_settlement_offer":
        return (
            "invoice_recovery_settlement_offer.docx",
            "Settlement offer DOCX",
            "Draft Settlement Offer",
            "This draft settlement offer packages the supported balance and evidence picture for an agreement-first recovery path.",
        )
    if route_status == "ready_for_payment_plan":
        return (
            "invoice_recovery_payment_plan_proposal.docx",
            "Payment plan proposal DOCX",
            "Draft Payment Plan Proposal",
            "This draft payment-plan proposal frames a structured repayment option using the supported balance and source record.",
        )
    if route_status == "ready_for_counsel_handoff":
        return (
            "invoice_recovery_counsel_handoff_memo.docx",
            "Counsel handoff memo DOCX",
            "Counsel Handoff Memo",
            "This memo summarizes the supported balance, current evidence, and handoff notes for downstream counsel review.",
        )
    if route_status == "ready_for_court_draft":
        return (
            "invoice_recovery_court_draft_memo.docx",
            "Court draft prep memo DOCX",
            "Court Draft Preparation Memo",
            "This memo summarizes the current litigation-ready record before complaint drafting or advanced court work begins.",
        )
    return (
        "invoice_recovery_demand_letter.docx",
        "Demand letter DOCX",
        "Draft Demand Letter",
        "This draft demand letter packages the supported balance and evidence picture for a pre-suit demand path.",
    )


def _materialize_invoice_recovery_outputs(
    record: CaseRecord,
    summary: dict[str, object],
    missing_documents: list[str],
    next_step: str,
    route_status: str,
    generation_allowed: bool,
) -> list[ActionPacketOutput]:
    output_dir = ensure_case_dirs(record.id)["output_dir"]
    created_at = datetime.now(timezone.utc).strftime("%B %d, %Y")

    balance_pdf = output_dir / "invoice_recovery_balance_summary.pdf"
    checklist_pdf = output_dir / "invoice_recovery_evidence_checklist.pdf"
    primary_output_name, primary_output_label, primary_output_title, primary_output_intro = _invoice_primary_output_spec(route_status)
    primary_docx = output_dir / primary_output_name

    _write_simple_pdf(
        balance_pdf,
        "DocketMint Invoice Recovery Balance Summary",
        [
            f"Date: {created_at}",
            f"Matter: {record.title or record.reference or record.id}",
            f"Supported balance: {_money(summary.get('supported_balance'))}",
            f"Invoices found: {summary.get('invoice_count') or 0}",
            f"Has contract: {'Yes' if summary.get('has_contract') else 'No'}",
            f"Has guaranty: {'Yes' if summary.get('has_personal_guaranty') else 'No'}",
            f"Has attorneys' fees clause: {'Yes' if summary.get('has_attorneys_fees_clause') else 'No'}",
            f"Has venue clause: {'Yes' if summary.get('has_venue_clause') else 'No'}",
            f"Has performance proof: {'Yes' if summary.get('performance_proof_present') else 'No'}",
            f"Recommended next step: {str(next_step).replace('_', ' ')}",
            "Human review remains important before sending a demand or filing.",
        ],
    )

    checklist_lines = [
        f"Date: {created_at}",
        f"Matter: {record.title or record.reference or record.id}",
        "Missing or weak evidence checklist:",
    ]
    if missing_documents:
        checklist_lines.extend([f"- {item}" for item in missing_documents])
    else:
        checklist_lines.append("- No major evidence gaps were flagged from the uploaded sources.")
    checklist_lines.append("Nothing is sent automatically.")
    _write_simple_pdf(checklist_pdf, "DocketMint Invoice Recovery Evidence Checklist", checklist_lines)

    if generation_allowed:
        doc = Document()
        doc.add_heading(primary_output_title, 0)
        doc.add_paragraph("DocketMint generated draft. Human review required before sending.")
        doc.add_paragraph(f"Date: {created_at}")
        doc.add_paragraph(f"Reference: {record.reference or record.id}")
        doc.add_paragraph("")
        doc.add_paragraph(primary_output_intro)
        doc.add_paragraph(
            f"Our current supported balance review indicates approximately {_money(summary.get('supported_balance'))} remains due. "
            f"The recommended next step is {str(next_step).strip().lower()}."
        )
        doc.add_paragraph(
            "Review the draft carefully before sending, exporting, or using it as a handoff artifact. "
            "Nothing is sent automatically."
        )
        if missing_documents:
            doc.add_paragraph(
                "Internal note before sending: the following support items should still be reviewed or attached if available: "
                + ", ".join(missing_documents) + "."
            )
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("DocketMint / VitaCoreX workflow draft")
        doc.save(primary_docx)

    outputs = [
        ActionPacketOutput(
            type="balance_summary_pdf",
            label="Balance summary PDF",
            status="generated",
            download_url=_case_file_url(record, balance_pdf),
            file_path=str(balance_pdf),
            owner_kind="case",
            owner_id=record.id,
            service_slug=record.workflow_slug or "invoice_recovery",
        ),
        ActionPacketOutput(
            type="evidence_gap_checklist_pdf",
            label="Evidence gap checklist PDF",
            status="generated",
            download_url=_case_file_url(record, checklist_pdf),
            file_path=str(checklist_pdf),
            owner_kind="case",
            owner_id=record.id,
            service_slug=record.workflow_slug or "invoice_recovery",
        ),
    ]
    if generation_allowed:
        outputs.append(
            ActionPacketOutput(
                type="recovery_primary_output_docx",
                label=primary_output_label,
                status="generated",
                download_url=_case_file_url(record, primary_docx),
                file_path=str(primary_docx),
                owner_kind="case",
                owner_id=record.id,
                service_slug=record.workflow_slug or "invoice_recovery",
            )
        )
    return outputs


def _fallback_irs_notice_payload(record: CaseRecord, payload: dict[str, object]) -> dict[str, object]:
    if payload.get("notice_number") and payload.get("recommended_route"):
        return payload

    texts: list[str] = []
    for item in active_source_files(record):
        text = str(item.extracted_text or "").strip()
        if text:
            texts.append(text)
            continue
        storage_path = Path(str(item.storage_path or ""))
        if storage_path.exists() and storage_path.suffix.lower() in {".txt", ".md"}:
            texts.append(storage_path.read_text(encoding="utf-8", errors="ignore"))
        else:
            texts.append(f"{item.original_filename} {item.file_type}")

    if not texts:
        return payload

    repaired = parse_irs_notice(texts)
    merged = dict(payload)
    for key, value in repaired.items():
        if value not in (None, "", [], {}):
            merged[key] = value
    workflow_payload = dict(record.workflow_payload or {})
    workflow_payload["irs_notice"] = merged
    record.workflow_payload = workflow_payload
    return merged


def _materialize_irs_notice_outputs(
    record: CaseRecord,
    summary: dict[str, object],
    missing_documents: list[str],
    next_step: str,
) -> list[ActionPacketOutput]:
    output_dir = ensure_case_dirs(record.id)["output_dir"]
    created_at = datetime.now(timezone.utc).strftime("%B %d, %Y")

    notice_pdf = output_dir / "irs_notice_summary.pdf"
    checklist_pdf = output_dir / "irs_response_checklist.pdf"
    response_docx = output_dir / "irs_response_letter.docx"

    _write_simple_pdf(
        notice_pdf,
        "DocketMint IRS Notice Summary",
        [
            f"Date: {created_at}",
            f"Matter: {record.title or record.reference or record.id}",
            f"Notice number: {summary.get('notice_number') or 'Not detected'}",
            f"Tax year: {summary.get('tax_year') or 'Not detected'}",
            f"Claimed amount: {_money(summary.get('claimed_amount')) if summary.get('claimed_amount') else '-'}",
            f"Recommended next step: {str(next_step).replace('_', ' ')}",
            "Keywords detected: " + (", ".join(summary.get("keywords") or []) or "None detected"),
            "Nothing is sent automatically. Review the notice and supporting records before responding.",
        ],
    )

    checklist_lines = [
        f"Date: {created_at}",
        f"Matter: {record.title or record.reference or record.id}",
        "What to gather before responding:",
    ]
    if missing_documents:
        checklist_lines.extend([f"- {item}" for item in missing_documents])
    else:
        checklist_lines.append("- No obvious missing documents were flagged from the uploaded notice set.")
    checklist_lines.append("Human review is still important before mailing, faxing, or uploading a response.")
    _write_simple_pdf(checklist_pdf, "DocketMint IRS Response Checklist", checklist_lines)

    doc = Document()
    doc.add_heading("Draft IRS Response Letter", 0)
    doc.add_paragraph("DocketMint generated draft. Human review required before sending.")
    doc.add_paragraph(f"Date: {created_at}")
    doc.add_paragraph(f"Notice number: {summary.get('notice_number') or 'Unknown'}")
    if summary.get("tax_year"):
        doc.add_paragraph(f"Tax year: {summary.get('tax_year')}")
    if summary.get("claimed_amount"):
        doc.add_paragraph(f"Claimed amount referenced in notice: {_money(summary.get('claimed_amount'))}")
    doc.add_paragraph("")
    doc.add_paragraph("To whom it may concern:")
    doc.add_paragraph(
        "I am writing in response to the IRS notice identified above. I am reviewing the matter and preparing the supporting "
        "records needed to confirm the correct response path."
    )
    doc.add_paragraph(
        "Please apply this response together with any supporting records that are attached after final review. If any deadline "
        "or requested record needs clarification, I request that the account reflect that this response is under review."
    )
    if missing_documents:
        doc.add_paragraph(
            "Internal review note before sending: gather or confirm the following support items if available: "
            + ", ".join(missing_documents) + "."
        )
    doc.add_paragraph("Nothing in this draft is sent automatically.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("DocketMint / VitaCoreX workflow draft")
    doc.save(response_docx)

    return [
        ActionPacketOutput(
            type="notice_summary_pdf",
            label="IRS notice summary PDF",
            status="generated",
            download_url=_case_file_url(record, notice_pdf),
            file_path=str(notice_pdf),
            owner_kind="case",
            owner_id=record.id,
            service_slug=record.workflow_slug or "irs_notice",
        ),
        ActionPacketOutput(
            type="response_checklist_pdf",
            label="IRS response checklist PDF",
            status="generated",
            download_url=_case_file_url(record, checklist_pdf),
            file_path=str(checklist_pdf),
            owner_kind="case",
            owner_id=record.id,
            service_slug=record.workflow_slug or "irs_notice",
        ),
        ActionPacketOutput(
            type="response_letter_docx",
            label="IRS response letter DOCX",
            status="generated",
            download_url=_case_file_url(record, response_docx),
            file_path=str(response_docx),
            owner_kind="case",
            owner_id=record.id,
            service_slug=record.workflow_slug or "irs_notice",
        ),
    ]


def build_action_packet(record: CaseRecord) -> dict[str, object]:
    refresh_workspace_modules(record)
    issue = str(record.issue_category or "unknown")
    readiness = assess_packet_readiness(record)
    route_decision = determine_recovery_route(record) if issue == "invoice_recovery_matter" else None
    source_files = active_source_files(record)
    missing_documents: list[str] = []
    outputs: list[ActionPacketOutput] = []

    if not source_files and issue != "invoice_recovery_matter":
        missing_documents.append("source_documents")

    next_step = (record.recommended_paths or [record.client_snapshot.next_action if record.client_snapshot else "review_matter"])[0]
    if route_decision:
        next_step = (route_decision.get("primary_action") or {}).get("label") or next_step
    summary: dict[str, object] = {
        "issue_category": issue,
        "workflow_slug": record.workflow_slug,
        "current_stage": record.current_stage or record.status,
        "urgency_level": record.urgency_level or "medium",
        "title": record.title or record.reference or record.id,
        "status": record.status,
    }

    if issue == "invoice_recovery_matter":
        payload = _fallback_invoice_payload(record, (record.workflow_payload or {}).get("invoice_recovery", {}))
        truth = dict(record.extracted_case_data.recovery_truth or {})
        summary.update(
            {
                "supported_balance": truth.get("supported_balance") or payload.get("supported_balance"),
                "unsupported_balance": truth.get("unsupported_balance"),
                "invoice_count": truth.get("invoice_count") or payload.get("invoice_count"),
                "matched_invoice_count": truth.get("matched_invoice_count"),
                "invoice_numbers": payload.get("invoice_numbers", []),
                "has_contract": payload.get("has_contract"),
                "has_personal_guaranty": payload.get("has_personal_guaranty"),
                "has_attorneys_fees_clause": payload.get("has_attorneys_fees_clause"),
                "has_governing_law_clause": payload.get("has_governing_law_clause"),
                "has_venue_clause": payload.get("has_venue_clause"),
                "performance_proof_present": payload.get("performance_proof_present"),
                "evidence_gap_count": truth.get("evidence_gap_count") or payload.get("evidence_gap_count"),
            }
        )
        missing_documents = list(route_decision.get("missing_evidence") or [])
        outputs.extend(
            _materialize_invoice_recovery_outputs(
                record,
                summary,
                missing_documents,
                next_step,
                str(route_decision.get("status") or "blocked_missing_evidence"),
                bool(route_decision.get("generation_allowed")),
            )
        )
    elif issue == "irs_notice":
        payload = _fallback_irs_notice_payload(record, (record.workflow_payload or {}).get("irs_notice", {}))
        summary.update(
            {
                "notice_number": payload.get("notice_number"),
                "tax_year": payload.get("tax_year"),
                "claimed_amount": payload.get("claimed_amount"),
                "keywords": payload.get("keywords", []),
                "recommended_route": payload.get("recommended_route"),
            }
        )
        if not source_files:
            missing_documents.append("notice_copy")
        keywords = [str(item).strip().lower() for item in (payload.get("keywords") or [])]
        if "payment" in keywords:
            missing_documents.append("payment_records_if_any")
        if "identity" in keywords:
            missing_documents.append("identity_verification_records")
        if "audit" in keywords or "correction" in keywords:
            missing_documents.append("supporting_explanation_statement")
        outputs.extend(_materialize_irs_notice_outputs(record, summary, list(dict.fromkeys(missing_documents)), next_step))
    elif issue == "business_contract":
        outputs.extend(
            [
                ActionPacketOutput(type="contract_risk_summary", label="Contract risk summary", status="available"),
                ActionPacketOutput(type="stronger_contract_draft", label="Stronger contract draft", status="available"),
                ActionPacketOutput(type="change_memo", label="Change memo", status="available"),
            ]
        )
    elif issue == "vehicle_purchase_contract":
        outputs.extend(
            [
                ActionPacketOutput(type="deal_audit_summary", label="Deal audit summary", status="available"),
                ActionPacketOutput(type="pricing_reference_sheet", label="Pricing reference sheet", status="available"),
                ActionPacketOutput(type="negotiation_guidance", label="Negotiation guidance", status="available"),
            ]
        )
    elif issue == "immigration_intake":
        outputs.extend(
            [
                ActionPacketOutput(type="mapped_form_pdf", label="Mapped form draft", status="available"),
                ActionPacketOutput(type="evidence_checklist", label="Evidence checklist", status="available"),
            ]
        )
    else:
        outputs.extend(
            [
                ActionPacketOutput(type="matter_summary", label="Matter summary", status="available"),
                ActionPacketOutput(type="next_step_card", label="Next step card", status="available"),
            ]
        )

    outputs.extend(_existing_outputs(record))
    deduped: list[ActionPacketOutput] = []
    seen: set[tuple[str, str | None]] = set()
    for item in outputs:
        key = (item.type, item.label)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    blockers = list(readiness.blockers or [])
    warnings = list(readiness.warnings or [])
    if route_decision:
        route_missing_blockers = (route_decision.get("missing_evidence") or []) if not route_decision.get("generation_allowed") else []
        blockers = list(
            dict.fromkeys(
                [
                    *(route_decision.get("blockers") or []),
                    *route_missing_blockers,
                    *blockers,
                ]
            )
        )

    return {
        "summary": summary,
        "next_step": next_step,
        "missing_documents": list(dict.fromkeys(missing_documents)),
        "missing_evidence": list(route_decision.get("missing_evidence") or []) if route_decision else list(dict.fromkeys(missing_documents)),
        "contacts": [item.model_dump(mode="json") for item in (record.official_contacts or [])],
        "deadlines": [item.model_dump(mode="json") for item in (record.deadlines or [])],
        "outputs": [item.model_dump(mode="json") for item in deduped],
        "blockers": blockers,
        "warnings": warnings,
        "route_status": route_decision.get("status") if route_decision else None,
        "route_status_label": route_decision.get("status_label") if route_decision else None,
        "route_reason_codes": list(route_decision.get("reason_codes") or []) if route_decision else [],
        "route_reasons": list(route_decision.get("reasons") or []) if route_decision else [],
        "primary_next_step": route_decision.get("primary_action") if route_decision else None,
        "optional_next_steps": list(route_decision.get("optional_actions") or []) if route_decision else [],
        "secondary_options": list(route_decision.get("secondary_options") or route_decision.get("optional_actions") or []) if route_decision else [],
        "readiness_score": route_decision.get("readiness_score") if route_decision else None,
        "source_backed_explanation": list(route_decision.get("source_backed_explanation") or []) if route_decision else [],
        "generation_allowed": bool(route_decision.get("generation_allowed")) if route_decision else readiness.is_ready,
        "reconciliation_center": build_reconciliation_center(record),
        "communications_last_touch": record.communications_log[0].model_dump(mode="json") if record.communications_log else None,
        "settlement_tracker": [item.model_dump(mode="json") for item in (record.settlement_tracker or [])[:3]],
        "packet_archive": [item.model_dump(mode="json") for item in (record.packet_archive or [])[:3]],
        "nothing_is_sent_automatically": True,
    }
