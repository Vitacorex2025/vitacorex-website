from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from backend.main import app
from backend.models.case_schema import CaseRecord, DebtCaseExtraction, PartyDetails, PartyRole
from backend.tests.auth_test_utils import authenticated_client
from backend.services.amount_parser import parse_amount
from backend.services.docx_generator import _qa_generated_document

client = authenticated_client(app)
DATA_DIR = Path(__file__).resolve().parents[1] / "sample_inputs"

XpressMixedFiles = [
    "Cases to Steven Citifuel.xlsx",
    "Cases to Steven TSS.xlsx",
    "XPRESS GLOBAL TRANSPORTATION.pdf",
    "XPRESS_GLOBAL_TRANSPORTATIONcom_Cube_Systems_LLCorg_Cube_Systems.pdf",
    "XPRESS_GLOBAL_TRANSPORTATIONcom_Cube_Systems_LLCorg_Cube_Systems (2).pdf",
    "XPRESS_GLOBAL_TRANSPORTATIONcom_Cube_Systems_LLCorg_Cube_Systems (3).pdf",
]


def _create_case(title: str, plaintiff: str | None = None, defendant: str | None = None) -> str:
    response = client.post(
        "/api/cases",
        json={"matter_name": title, "plaintiff_name": plaintiff, "defendant_name": defendant},
    )
    assert response.status_code == 200
    return response.json()["id"]


def _upload_files(case_id: str, filenames: list[str]) -> None:
    files = []
    handles = []
    try:
        for name in filenames:
            path = DATA_DIR / name
            handle = path.open("rb")
            handles.append(handle)
            content_type = "application/pdf" if path.suffix.lower() == ".pdf" else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            files.append(("files", (name, handle, content_type)))
        response = client.post(f"/api/cases/{case_id}/source-files", files=files)
        assert response.status_code == 200
    finally:
        for handle in handles:
            handle.close()


def _process_case(case_id: str) -> dict:
    response = client.post(f"/api/cases/{case_id}/process")
    assert response.status_code == 200
    return response.json()


def _create_xpress_mixed_case() -> tuple[str, dict]:
    case_id = _create_case("Xpress mixed debtor upload")
    _upload_files(case_id, XpressMixedFiles)
    return case_id, _process_case(case_id)


def _find_xpress_child(created_cases: list[dict[str, str]]) -> str:
    for item in created_cases:
        if "XPRESS" in item["debtor_name"].upper():
            return item["case_id"]
    raise AssertionError("XPRESS child case was not created.")


def _split_to_xpress_child(case_id: str) -> tuple[str, dict]:
    split = client.post(f"/api/cases/{case_id}/split-by-debtor")
    assert split.status_code == 200
    child_case_id = _find_xpress_child(split.json()["created_cases"])
    processed_child = _process_case(child_case_id)
    return child_case_id, processed_child


def _approve_xpress_child(case_id: str) -> dict:
    review = client.patch(
        f"/api/cases/{case_id}/review",
        json={"venue_county": "Orange", "governing_law": "Florida"},
    )
    assert review.status_code == 200
    return review.json()


def _build_invoice_workbook(path: Path, company: str, amount_text: str, organization: str = "Cube Systems LLC") -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    sheet.append(["Organization", "Company", "Invoice Number", "Invoice Date", "Remaining Amount"])
    sheet.append([organization, company, "INV-100", "2025-02-20", amount_text])
    workbook.save(path)


def _build_invoice_pdf(path: Path, company: str, total_amount: str, plaintiff: str = "Cube Systems LLC") -> None:
    pdf = canvas.Canvas(str(path))
    y = 780
    for line in [
        plaintiff,
        "Address: 2025 Riverside Dr Ste 323 Columbus OH",
        "Invoice INV-100",
        "Invoice Date 02/20/2025",
        "Due Date 02/21/2025",
        f"Total amount due $ {total_amount}",
        company,
        "Address: 3910 auburn blvd apt 11 Sacramento California",
    ]:
        pdf.drawString(72, y, line)
        y -= 18
    pdf.save()


def test_mixed_debtor_files_in_one_case_are_blocked():
    case_id, processed = _create_xpress_mixed_case()

    assert processed["status"] == "blocked"
    extraction = processed["extracted_case_data"]
    consistency = extraction["matter_consistency"]
    reconciliation = extraction["amount_reconciliation"]

    assert consistency["status"] == "blocked"
    assert "This upload contains multiple debtor matters. Split them into separate cases." in consistency["blockers"]
    assert "XPRESS GLOBAL TRANSPORTATION" in consistency["detected_debtor_entities"]
    assert len(consistency["detected_debtor_entities"]) > 1
    assert extraction["plaintiff"]["party_name"] == "Cube Systems LLC"
    assert extraction["defendant_entity"]["party_name"] == "XPRESS GLOBAL TRANSPORTATION"
    assert extraction["personal_guarantor"]["party_name"] == "Kambuljon Kazakov"
    assert extraction["amount_in_controversy"] == "110651.07"
    assert reconciliation["status"] == "pass"
    assert reconciliation["source_spreadsheet_total"] == "110651.07"
    assert reconciliation["invoice_pdf_total"] == "110651.07"
    assert client.post(f"/api/generate/{case_id}?wait=true", json={"selected_document_types": ["complaint"]}).status_code == 400


def test_multi_debtor_workbook_is_blocked_until_split_by_debtor():
    case_id = _create_case("Citifuel workbook only")
    _upload_files(case_id, ["Cases to Steven Citifuel.xlsx"])
    processed = _process_case(case_id)
    consistency = processed["extracted_case_data"]["matter_consistency"]

    assert consistency["status"] == "blocked"
    assert any("multiple debtor companies" in message.lower() for message in consistency["blockers"])

    split = client.post(f"/api/cases/{case_id}/split-by-debtor")
    assert split.status_code == 200
    created_cases = split.json()["created_cases"]
    assert len(created_cases) >= 2
    assert any("XPRESS" in item["debtor_name"].upper() for item in created_cases)
    assert any("AFRUZ" in item["debtor_name"].upper() for item in created_cases)


def test_malformed_numeric_strings_do_not_inflate_balances_silently():
    parsed = parse_amount("9,301,33")
    assert parsed.value is None
    assert parsed.confidence == 0.0
    assert "ambiguous" in (parsed.warning or "").lower()

    parsed_second = parse_amount("13,989,97")
    assert parsed_second.value is None
    assert parsed_second.confidence == 0.0
    assert "ambiguous" in (parsed_second.warning or "").lower()


def test_xpress_guaranty_language_is_detected_and_clean_subset_generation_works():
    parent_case_id, _ = _create_xpress_mixed_case()
    child_case_id, processed_child = _split_to_xpress_child(parent_case_id)

    extraction = processed_child["extracted_case_data"]
    assert processed_child["status"] == "ready_for_verification"
    assert extraction["matter_consistency"]["status"] == "pass"
    assert extraction["personal_guarantor"]["party_name"] == "Kambuljon Kazakov"
    guaranty_count = next(count for count in extraction["legal_counts"] if count["code"] == "breach_of_personal_guaranty")
    assert guaranty_count["supported"] is True

    _approve_xpress_child(child_case_id)
    generate = client.post(
        f"/api/generate/{child_case_id}?wait=true",
        json={"selected_document_types": ["complaint", "debt_calculation_sheet"]},
    )
    assert generate.status_code == 200
    payload = generate.json()
    generated_docs = payload["generated_documents"]
    assert [doc["document_type"] for doc in generated_docs] == ["complaint", "debt_calculation_sheet"]
    assert all(doc["status"] == "generated" for doc in generated_docs)

    complaint_doc = next(doc for doc in generated_docs if doc["document_type"] == "complaint")
    response = client.get(complaint_doc["download_url"])
    assert response.status_code == 200
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
        temp_docx.write(response.content)
        temp_docx.flush()
        temp_docx_path = temp_docx.name
    try:
        text = "\n".join(paragraph.text for paragraph in Document(temp_docx_path).paragraphs)
    finally:
        Path(temp_docx_path).unlink(missing_ok=True)

    assert "Kambuljon Kazakov" in text
    assert text.count("WHEREFORE") == 1
    assert "Defendant  personally" not in text


def test_amount_reconciliation_mismatch_blocks_generation(tmp_path: Path):
    case_id = _create_case(
        "Synthetic mismatch",
        plaintiff="Cube Systems LLC",
        defendant="Single Matter LLC",
    )

    workbook_path = tmp_path / "synthetic.xlsx"
    pdf_path = tmp_path / "synthetic invoice.pdf"
    _build_invoice_workbook(workbook_path, "Single Matter LLC", "100.00")
    _build_invoice_pdf(pdf_path, "Single Matter LLC", "120.00")

    with workbook_path.open("rb") as workbook_fh, pdf_path.open("rb") as pdf_fh:
        upload = client.post(
            f"/api/cases/{case_id}/source-files",
            files=[
                ("files", (workbook_path.name, workbook_fh, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
                ("files", (pdf_path.name, pdf_fh, "application/pdf")),
            ],
        )
    assert upload.status_code == 200

    processed = _process_case(case_id)
    extraction = processed["extracted_case_data"]
    reconciliation = extraction["amount_reconciliation"]

    assert reconciliation["status"] == "blocked"
    assert "120.00" in reconciliation["blockers"][0]
    assert "100.00" in reconciliation["blockers"][0]

    client.patch(f"/api/cases/{case_id}/review", json={"venue_county": "Orange", "governing_law": "Florida"})
    generate = client.post(
        f"/api/generate/{case_id}?wait=true",
        json={"selected_document_types": ["debt_calculation_sheet"]},
    )
    assert generate.status_code == 400
    payload = generate.json()
    assert "amount_reconciliation" in str(payload)
    assert "120.00" in str(payload)
    assert "100.00" in str(payload)


def test_stale_case_summary_amount_is_replaced_after_source_cleanup():
    case_id, _ = _create_xpress_mixed_case()
    mixed_list_entry = next(item for item in client.get("/api/cases?limit=20").json() if item["id"] == case_id)
    assert mixed_list_entry["amount_in_controversy"] == "110651.07"

    current_case = client.get(f"/api/cases/{case_id}").json()
    for source in current_case["source_files"]:
        name = (source.get("original_filename") or "").upper()
        keep = name.startswith("XPRESS")
        if not keep:
            client.patch(
                f"/api/cases/{case_id}/source-files/{source['id']}",
                json={"active": False, "reason": "Removed from Xpress-only active matter."},
            )

    processed = _process_case(case_id)
    assert processed["status"] == "ready_for_verification"
    assert processed["extracted_case_data"]["defendant_entity"]["party_name"] == "XPRESS GLOBAL TRANSPORTATION"
    listed = next(item for item in client.get("/api/cases?limit=20").json() if item["id"] == case_id)
    assert listed["amount_in_controversy"] == "110651.07"


def test_upload_scope_patch_excludes_multi_debtor_workbook_from_active_matter():
    case_id = _create_case("Upload scope toggle")
    _upload_files(case_id, ["Cases to Steven Citifuel.xlsx", "XPRESS GLOBAL TRANSPORTATION.pdf"])
    case_payload = client.get(f"/api/cases/{case_id}").json()
    workbook = next(item for item in case_payload["source_files"] if item["original_filename"] == "Cases to Steven Citifuel.xlsx")

    patched = client.patch(
        f"/api/cases/{case_id}/source-files/{workbook['id']}",
        json={"active": False, "reason": "Exclude multi-debtor workbook from single-matter processing."},
    )
    assert patched.status_code == 200
    refreshed = client.get(f"/api/cases/{case_id}").json()
    workbook_refreshed = next(item for item in refreshed["source_files"] if item["id"] == workbook["id"])
    assert workbook_refreshed["excluded_from_active_matter"] is True
    assert refreshed["status"] == "uploaded"


def test_blocked_case_cannot_remain_packet_generated_after_new_contaminating_upload():
    parent_case_id, _ = _create_xpress_mixed_case()
    child_case_id, _ = _split_to_xpress_child(parent_case_id)
    _approve_xpress_child(child_case_id)
    generated = client.post(
        f"/api/generate/{child_case_id}?wait=true",
        json={"selected_document_types": ["complaint"]},
    )
    assert generated.status_code == 200
    case_after_generation = client.get(f"/api/cases/{child_case_id}").json()
    assert case_after_generation["status"] == "packet_generated"
    assert case_after_generation["packet_generation"] is not None

    _upload_files(child_case_id, ["Cases to Steven Citifuel.xlsx"])
    contaminated = _process_case(child_case_id)
    assert contaminated["status"] == "blocked"
    assert contaminated["packet_generation"] is None


def test_sunbiz_no_match_and_verified_branches_are_distinct(monkeypatch):
    from backend.api.routes import sunbiz as sunbiz_route
    from backend.services.sunbiz_lookup import SunbizLookupPayload

    case_id = _create_case("SunBiz branch coverage", plaintiff="Cube Systems LLC", defendant="Test Entity LLC")

    def fake_no_match(name: str):
        return SunbizLookupPayload(
            queried_entity_name=name,
            matched=False,
            exact_match=False,
            match_score=0.11,
            warnings=["No Florida Division of Corporations match found for the queried entity."],
            provider_status="ok",
            result_classification="no_match",
            official_record_available=False,
            official_source_urls=["https://search.sunbiz.org/inquiry/corporationsearch/byname"],
        )

    monkeypatch.setattr(sunbiz_route, "lookup_entity", fake_no_match)
    no_match = client.post(f"/api/sunbiz/cases/{case_id}/lookup")
    assert no_match.status_code == 200
    assert no_match.json()["result_classification"] == "no_match"
    assert no_match.json()["registered_agent_name"] is None

    def fake_verified(name: str):
        return SunbizLookupPayload(
            queried_entity_name=name,
            matched=True,
            exact_match=True,
            match_score=1.0,
            matched_entity_name="TEST ENTITY LLC",
            detail_url="https://search.sunbiz.org/Inquiry/corporationsearch/SearchResultDetail?x=1",
            florida_document_number="L24000012345",
            fein_number="12-3456789",
            entity_status="ACTIVE",
            entity_type="Florida Limited Liability Company",
            principal_address="123 MAIN ST, ORLANDO, FL 32801",
            registered_agent_name="REGISTERED AGENT LLC",
            registered_agent_address="123 MAIN ST, ORLANDO, FL 32801",
            officers=[{"officer_name": "Jane Manager", "officer_title": "MGR", "officer_address": "123 MAIN ST"}],
            provider_status="ok",
            result_classification="verified_match",
            official_record_available=True,
            official_source_urls=["https://search.sunbiz.org/inquiry/corporationsearch/byname"],
            source_url="https://search.sunbiz.org/Inquiry/corporationsearch/SearchResultDetail?x=1",
        )

    monkeypatch.setattr(sunbiz_route, "lookup_entity", fake_verified)
    verified = client.post(f"/api/sunbiz/cases/{case_id}/lookup")
    assert verified.status_code == 200
    payload = verified.json()
    assert payload["result_classification"] == "verified_match"
    assert payload["registered_agent_name"] == "REGISTERED AGENT LLC"
    assert payload["official_record_available"] is True


def test_document_qa_detects_duplicate_wherefore_and_blank_placeholders(tmp_path: Path):
    complaint_path = tmp_path / "bad complaint.docx"
    document = Document()
    document.add_paragraph("Defendant  personally signed and/or guaranteed the account obligations.")
    document.add_paragraph("WHEREFORE, Plaintiff requests relief.")
    document.add_paragraph("WHEREFORE, Plaintiff requests relief again.")
    document.add_paragraph("Balance due is $1,000.00.")
    document.save(complaint_path)

    record = CaseRecord(
        id="case_qa",
        extracted_case_data=DebtCaseExtraction(
            plaintiff=PartyDetails(party_name="Cube Systems LLC", party_role=PartyRole.PLAINTIFF),
            defendant_entity=PartyDetails(party_name="Single Matter LLC", party_role=PartyRole.DEFENDANT_ENTITY),
            personal_guarantor=PartyDetails(party_name="Jane Doe", party_role=PartyRole.PERSONAL_GUARANTOR),
            amount_in_controversy="1000.00",
        ),
    )
    record.extracted_case_data.amount_reconciliation.approved_total = "1000.00"

    issues = _qa_generated_document(record, "complaint", complaint_path)
    assert any("WHEREFORE" in issue for issue in issues)
    assert any("blank defendant" in issue.lower() for issue in issues)


def test_demand_letter_without_address_is_blocked_before_packet_generation():
    case_id = _create_case(
        "Counterparty demand letter readiness",
        plaintiff="Client Company LLC",
        defendant="Counterparty LLC",
    )
    _upload_files(case_id, ["Counterparty LLC.pdf", "Counterparty LLC Invoices.xlsx"])
    _process_case(case_id)

    registry = client.get(f"/api/templates/registry?case_id={case_id}")
    assert registry.status_code == 200
    payload = registry.json()

    demand_letter = next(item for item in payload["templates"] if item["document_type"] == "demand_letter")
    assert demand_letter["generation_allowed"] is False
    assert any("address" in message.lower() for message in demand_letter["blocking_reasons"])

    allowed_document_types = [item["document_type"] for item in payload["templates"] if item["generation_allowed"]]
    assert "demand_letter" not in allowed_document_types

    generate = client.post(
        f"/api/generate/{case_id}?wait=true",
        json={"selected_document_types": allowed_document_types},
    )
    assert generate.status_code == 200
    generation_payload = generate.json()
    assert generation_payload["generation_progress"]["status"] == "completed"
    assert generation_payload["generation_progress"]["message"] == "6 generated / 6 requested"
    generated_types = [item["document_type"] for item in generation_payload["generated_documents"]]
    assert "demand_letter" not in generated_types


def test_reconciliation_center_persists_summary_state_after_processing():
    case_id, processed = _create_xpress_mixed_case()
    modules = processed["workflow_payload"]["recovery_workspace_modules"]
    reconciliation_center = modules["reconciliation_center"]

    assert reconciliation_center["supported_balance"] == "110651.07"
    assert reconciliation_center["debtor_count"] >= 2
    assert reconciliation_center["invoice_count"] >= 1
    assert reconciliation_center["matched_invoice_count"] >= 1
    assert reconciliation_center["debtor_conflicts"]
    assert reconciliation_center["evidence_gaps"]
    assert any(item["file_name"] == "Cases to Steven Citifuel.xlsx" for item in reconciliation_center["source_balances_by_file"])

    fetched = client.get(f"/api/cases/{case_id}")
    assert fetched.status_code == 200
    fetched_center = fetched.json()["workflow_payload"]["recovery_workspace_modules"]["reconciliation_center"]
    assert fetched_center["supported_balance"] == reconciliation_center["supported_balance"]


def test_communications_and_settlement_tracker_are_saved_and_retrieved():
    case_id = _create_case("Layer 4 workspace state")

    note = client.post(
        f"/api/cases/{case_id}/communications",
        json={
            "entry_type": "manual_note",
            "channel": "phone_call",
            "direction": "outbound",
            "summary": "Called AP manager",
            "body": "Confirmed supporting guaranty will be sent tomorrow.",
        },
    )
    assert note.status_code == 200, note.text
    note_payload = note.json()["entry"]
    assert note_payload["summary"] == "Called AP manager"
    assert note_payload["actor_label"]

    settlement = client.post(
        f"/api/cases/{case_id}/settlement-tracker",
        json={
            "proposal_type": "payment_plan",
            "status": "proposed",
            "title": "Three-payment plan",
            "summary": "Counterparty asked for three monthly installments.",
            "term_summary": "3 monthly payments of $4,000 starting 2026-04-01.",
            "proposed_amount": "12000.00",
            "proposed_payment_count": 3,
            "proposed_first_payment_date": "2026-04-01",
            "linked_route_status": "ready_for_payment_plan",
            "reason_codes": ["payment_plan_requested"],
        },
    )
    assert settlement.status_code == 200, settlement.text
    tracker_entry = settlement.json()["entry"]
    assert tracker_entry["status"] == "proposed"
    assert tracker_entry["proposal_type"] == "payment_plan"

    communications = client.get(f"/api/cases/{case_id}/communications")
    assert communications.status_code == 200
    assert communications.json()["entries"][0]["summary"] == "Called AP manager"

    tracker = client.get(f"/api/cases/{case_id}/settlement-tracker")
    assert tracker.status_code == 200
    assert tracker.json()["entries"][0]["title"] == "Three-payment plan"

    case_payload = client.get(f"/api/cases/{case_id}").json()
    modules = case_payload["workflow_payload"]["recovery_workspace_modules"]
    assert modules["communications_log"]["count"] == 1
    assert modules["settlement_tracker"]["count"] == 1
    assert modules["settlement_tracker"]["latest"]["linked_route_status"] == "ready_for_payment_plan"


def test_deadline_overrides_are_auditable_and_merge_into_case_deadlines():
    case_id = _create_case("Deadline override workspace")

    override = client.post(
        f"/api/cases/{case_id}/deadlines/override",
        json={
            "deadline_type": "response",
            "label": "Counsel response due",
            "date": "2026-04-15",
            "reason": "Confirmed by attorney review.",
        },
    )
    assert override.status_code == 200, override.text
    payload = override.json()
    assert payload["override"]["actor_label"]

    deadlines = client.get(f"/api/cases/{case_id}/deadlines")
    assert deadlines.status_code == 200
    assert deadlines.json()["overrides"][0]["label"] == "Counsel response due"
    merged = deadlines.json()["deadlines"][0]
    assert merged["label"] == "Counsel response due"
    assert merged["source_kind"] == "manual_override"
    assert merged["override_reason"] == "Confirmed by attorney review."

    case_payload = client.get(f"/api/cases/{case_id}").json()
    assert case_payload["deadline_overrides"][0]["label"] == "Counsel response due"


def test_packet_archive_versions_increment_and_court_drafts_stay_review_gated():
    parent_case_id, _ = _create_xpress_mixed_case()
    child_case_id, _ = _split_to_xpress_child(parent_case_id)
    _approve_xpress_child(child_case_id)

    first_generate = client.post(
        f"/api/generate/{child_case_id}?wait=true",
        json={"selected_document_types": ["complaint", "debt_calculation_sheet"]},
    )
    assert first_generate.status_code == 200, first_generate.text
    after_first = client.get(f"/api/cases/{child_case_id}").json()
    first_archive = after_first["packet_archive"][0]
    assert first_archive["version_number"] == 1
    assert first_archive["release_state"] == "ready_for_review"
    assert first_archive["review_required"] is True
    assert first_archive["is_court_draft"] is True
    assert any(item["entry_type"] == "generated_draft" for item in after_first["communications_log"])

    second_generate = client.post(
        f"/api/generate/{child_case_id}?wait=true",
        json={"selected_document_types": ["debt_calculation_sheet"]},
    )
    assert second_generate.status_code == 200, second_generate.text
    after_second = client.get(f"/api/cases/{child_case_id}").json()
    versions = [item["version_number"] for item in after_second["packet_archive"][:2]]
    assert versions == [2, 1]

    downloads = client.get(f"/api/downloads/cases/{child_case_id}")
    assert downloads.status_code == 200
    assert downloads.json()["packet_archive"][0]["version_number"] == 2
