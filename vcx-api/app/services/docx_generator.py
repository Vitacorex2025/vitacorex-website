"""DOCX document generator — creates contract documents from templates.

Phase 8: Contract Generator — uses python-docx to build professional
DOCX files from the clause library in contract_templates.py.

All generated documents include a prominent legal disclaimer and
"DRAFT — NOT LEGAL ADVICE" watermark header.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger("vcx.docx_generator")

# python-docx imports — guarded for environments where not installed
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX generation unavailable.")

from .contract_templates import (
    CONTRACT_TYPES,
    DISCLAIMER_TEXT,
    build_contract_sections,
)


def is_available() -> bool:
    """Check if DOCX generation is available (python-docx installed)."""
    return _DOCX_AVAILABLE


def generate_contract_docx(
    contract_type: str,
    params: dict,
    include_disclaimer: bool = True,
) -> bytes:
    """Generate a DOCX contract document and return as bytes.

    Args:
        contract_type: Key from CONTRACT_TYPES (e.g., "nda", "service")
        params: Questionnaire answers dict
        include_disclaimer: Whether to include the legal disclaimer page

    Returns:
        DOCX file content as bytes

    Raises:
        RuntimeError: If python-docx is not installed
        ValueError: If contract_type is unknown
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    spec = CONTRACT_TYPES.get(contract_type)
    if not spec:
        raise ValueError(f"Unknown contract type: {contract_type}")

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ── Default font ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # ── Paragraph spacing for readability ──────────────────────────
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Draft watermark header ──────────────────────────────────────
    _add_draft_header(doc)

    # ── Disclaimer page (if enabled) ────────────────────────────────
    if include_disclaimer:
        _add_disclaimer_page(doc)

    # ── Document properties (Word metadata) ─────────────────────────
    doc.core_properties.title = spec["label"]
    doc.core_properties.author = "VitaCoreX Contract Generator"
    doc.core_properties.comments = "Draft document — not legal advice"

    # ── Title ───────────────────────────────────────────────────────
    title = doc.add_heading(spec["label"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)  # VCX brand navy
    title.paragraph_format.space_after = Pt(24)

    # ── Build sections from template ────────────────────────────────
    sections = build_contract_sections(contract_type, params)

    section_number = 0
    for heading_text, body_text in sections:
        # Special handling for preamble and signatures (no numbering)
        if heading_text in ("Agreement", "Recitals", "Signatures"):
            _add_section_unnumbered(doc, heading_text, body_text)
        else:
            section_number += 1
            _add_section_numbered(doc, section_number, heading_text, body_text)

    # ── Footer disclaimer ───────────────────────────────────────────
    _add_footer_disclaimer(doc)

    # ── Serialize to bytes ──────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_draft_header(doc: Document) -> None:
    """Add a 'DRAFT — NOT LEGAL ADVICE' header to all pages."""
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("DRAFT — NOT LEGAL ADVICE — FOR REFERENCE ONLY")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    run.font.bold = True


def _add_disclaimer_page(doc: Document) -> None:
    """Add a full disclaimer page before the contract content."""
    heading = doc.add_heading("Important Legal Notice", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

    disclaimer_para = doc.add_paragraph()
    disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = disclaimer_para.add_run(DISCLAIMER_TEXT)
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Generated by VitaCoreX Contract Generator\n"
        "This document should be reviewed by licensed counsel before use."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page break after disclaimer
    doc.add_page_break()


def _add_section_unnumbered(doc: Document, heading: str, body: str) -> None:
    """Add an unnumbered section (preamble, recitals, signatures)."""
    h = doc.add_heading(heading, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)

    # Split body by double newline into paragraphs
    for chunk in body.split("\n\n"):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        # Handle single newlines within a chunk as line breaks
        lines = chunk.split("\n")
        for i, line in enumerate(lines):
            run = para.add_run(line)
            run.font.size = Pt(11)
            if i < len(lines) - 1:
                run.add_break()


def _add_section_numbered(
    doc: Document,
    number: int,
    heading: str,
    body: str,
) -> None:
    """Add a numbered section (Article N: Heading)."""
    h = doc.add_heading(f"Article {number}. {heading}", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)

    for chunk in body.split("\n\n"):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        lines = chunk.split("\n")
        for i, line in enumerate(lines):
            run = para.add_run(line)
            run.font.size = Pt(11)
            if i < len(lines) - 1:
                run.add_break()


def _add_footer_disclaimer(doc: Document) -> None:
    """Add a small disclaimer to the page footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        "Generated by VitaCoreX — Not legal advice — Consult licensed counsel"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def generate_ai_contract_docx(
    ai_text: str,
    title: str = "Improved Contract",
) -> bytes:
    """Generate a DOCX from AI-generated contract text (markdown-like).

    Parses ## headings as section headers and body text as paragraphs.
    Returns DOCX file content as bytes.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed.")

    import re as _re

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Draft header
    _add_draft_header(doc)

    # Disclaimer page
    _add_disclaimer_page(doc)

    # Document properties
    doc.core_properties.title = title
    doc.core_properties.author = "VitaCoreX AI Contract Generator"
    doc.core_properties.comments = "AI-generated draft — not legal advice"

    # Title
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
    t.paragraph_format.space_after = Pt(24)

    # AI badge
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = badge.add_run("Generated with AI assistance by VitaCoreX Contract Intelligence")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2D, 0x8A, 0x82)
    r.font.italic = True
    badge.paragraph_format.space_after = Pt(18)

    # Parse AI output into sections
    lines = ai_text.split("\n")
    article_num = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # ## HEADING or ## ARTICLE N: HEADING
        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            # Check if it's a named section (AGREEMENT, RECITALS, SIGNATURES)
            upper = heading_text.upper()
            if upper in ("AGREEMENT", "RECITALS", "SIGNATURES"):
                h = doc.add_heading(heading_text, level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(8)
            else:
                # Article heading
                # Remove "ARTICLE N:" prefix if AI included it
                clean = _re.sub(
                    r'^ARTICLE\s+\d+\s*[:.]\s*',
                    '', heading_text, flags=_re.IGNORECASE
                )
                article_num += 1
                h = doc.add_heading(
                    f"Article {article_num}. {clean}", level=2
                )
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(8)
        elif stripped.startswith("# "):
            # Top-level heading (contract title) — skip, we already added one
            pass
        elif stripped.startswith("### "):
            # Sub-heading
            sub = stripped[4:].strip()
            h = doc.add_heading(sub, level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x4F, 0x46)
                run.font.size = Pt(12)
        else:
            # Body text
            # Handle markdown bold
            clean_text = _re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(6)

            # Check if it's a numbered item or bullet
            if _re.match(r'^\d+\.?\s', clean_text):
                para.paragraph_format.left_indent = Inches(0.25)
            elif clean_text.startswith(("- ", "• ", "* ")):
                clean_text = clean_text[2:]
                para.paragraph_format.left_indent = Inches(0.25)
                run = para.add_run("• ")
                run.font.size = Pt(11)

            # Handle (a), (b) style sub-items
            if _re.match(r'^\([a-z]\)', clean_text):
                para.paragraph_format.left_indent = Inches(0.5)

            run = para.add_run(clean_text)
            run.font.size = Pt(11)

            # Bold for signature lines
            if "________________________" in clean_text:
                run.font.size = Pt(11)
            elif clean_text.upper() == clean_text and len(clean_text) > 10:
                run.font.bold = True
                run.font.size = Pt(10)

    # Footer disclaimer
    _add_footer_disclaimer(doc)

    # Serialize
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_review_comments_docx(
    ai_text: str,
    original_filename: str = "contract",
) -> bytes:
    """Generate a DOCX with AI review comments for negotiation.

    Parses the AI output (### headings, **bold** fields) into a
    professional review document with color-coded priority sections.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed.")

    import re as _re

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Header
    header_section = doc.sections[0]
    header = header_section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("CONFIDENTIAL — CONTRACT REVIEW NOTES — FOR NEGOTIATION USE ONLY")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x2D, 0x8A, 0x82)
    run.font.bold = True

    # Title
    title = doc.add_heading("Contract Review Comments", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
    title.paragraph_format.space_after = Pt(8)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Review of: {original_filename}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.italic = True
    sub.paragraph_format.space_after = Pt(4)

    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = badge.add_run("Generated by VitaCoreX AI Contract Intelligence")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2D, 0x8A, 0x82)
    badge.paragraph_format.space_after = Pt(18)

    # Disclaimer
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = disc.add_run(
        "DISCLAIMER: These review comments are generated by an automated "
        "system and are provided as negotiation preparation notes only. "
        "They do NOT constitute legal advice. Consult with a licensed "
        "attorney before making any decisions or changes based on these comments."
    )
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    disc.paragraph_format.space_after = Pt(18)

    # Priority colors
    _priority_colors = {
        "HIGH": RGBColor(0x8B, 0x43, 0x48),
        "MEDIUM": RGBColor(0x9A, 0x6A, 0x20),
        "LOW": RGBColor(0x2F, 0x6B, 0x57),
    }

    # Parse AI output
    lines = ai_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # ### Section heading
        if stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            h = doc.add_heading(heading_text, level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            # Add a thin separator line
            sep = doc.add_paragraph()
            sep.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            h = doc.add_heading(heading_text, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x17, 0x3A, 0x63)
        elif stripped.startswith("# "):
            pass  # skip top-level titles
        else:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)

            # Handle **bold** sections
            clean = stripped
            # Check for priority tags
            priority_color = None
            for prio, color in _priority_colors.items():
                if prio in clean.upper():
                    priority_color = color
                    break

            # Parse **bold** and regular text
            parts = _re.split(r'(\*\*[^*]+\*\*)', clean)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    label = part[2:-2]
                    run = para.add_run(label)
                    run.font.bold = True
                    run.font.size = Pt(11)
                    if priority_color and ("Priority" in label or label in ("HIGH", "MEDIUM", "LOW")):
                        run.font.color.rgb = priority_color
                else:
                    run = para.add_run(part)
                    run.font.size = Pt(11)

            # Indent bullets and sub-items
            if clean.startswith(("- ", "• ", "* ")):
                para.paragraph_format.left_indent = Inches(0.25)

    # Footer
    footer = doc.sections[0].footer
    fpara = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fpara.add_run(
        "VitaCoreX Contract Intelligence — Confidential Review Notes — Not Legal Advice"
    )
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Serialize
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
