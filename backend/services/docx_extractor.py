
from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx_text(path: str | Path) -> str:
    document = Document(str(Path(path)))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join((paragraph.text or "").split()).strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(" ".join((cell.text or "").split()).strip() for cell in row.cells if cell.text)
            row_text = row_text.strip(" |")
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()
