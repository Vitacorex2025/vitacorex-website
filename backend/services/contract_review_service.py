from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from backend.models.case_schema import CaseRecord, ContractClauseFinding, ContractReviewResult
from backend.services.case_store import ensure_case_dirs


def _utcnow():
    return datetime.now(timezone.utc)


def _clean(value: object | None) -> str | None:
    text = " ".join(str(value or "").replace("\u00a0", " ").split()).strip(" \n\t\r,;")
    return text or None


def _contract_files(record: CaseRecord):
    return [item for item in record.source_files if str(item.file_type) in {"contract_pdf", "guaranty_pdf", "unknown"}]


def _combined_contract_text(record: CaseRecord) -> list[tuple[str, str]]:
    files = []
    for item in _contract_files(record):
        text = item.extracted_text or ""
        if text.strip():
            files.append((item.original_filename, text))
    return files


CLAUSE_LIBRARY: list[dict[str, Any]] = [
    {
        "clause_type": "payment_terms",
        "title": "Payment terms",
        "patterns": [r"((?:payment terms|invoice due|net\s+\d+|due upon receipt|payment due)[\s\S]{0,260})"],
        "essential": True,
        "strong_terms": ["net", "late fee", "suspend", "stop work", "immediately due"],
        "weak_terms": ["reasonable time", "promptly"],
        "dangerous_terms": [],
        "why": "Payment timing drives cash flow and decides when a default actually starts.",
        "wrong": "If the due date is vague, the other side can delay payment and still argue there was no default.",
        "improve": "State when payment is due, what happens after non-payment, and whether services can be suspended.",
        "fallback": (
            "Client shall pay each invoice within fifteen (15) days after receipt. Any unpaid amount accrues a late "
            "charge of the lesser of 1.5% per month or the maximum lawful rate. Provider may suspend further work "
            "after written notice if an invoice remains unpaid for five (5) business days after its due date."
        ),
        "rationale": "Detailed payment mechanics reduce collection ambiguity and support a cleaner demand record.",
        "jurisdiction_note": "Late-fee and interest limits can be state-specific and should be checked before final use.",
    },
    {
        "clause_type": "default_triggers",
        "title": "Default and acceleration",
        "patterns": [r"((?:default|event of default|failure to pay|material breach)[\s\S]{0,260})"],
        "essential": True,
        "strong_terms": ["event of default", "accelerate", "cure period", "suspend"],
        "weak_terms": ["material breach"],
        "dangerous_terms": [],
        "why": "A default clause converts non-performance into a defined trigger for remedies.",
        "wrong": "Without clear triggers, the breaching party can argue the contract never defined when remedies became available.",
        "improve": "Define payment default, cure periods, and whether remaining amounts can accelerate.",
        "fallback": (
            "Each of the following constitutes an Event of Default: failure to pay any invoice when due, breach of "
            "a material covenant that is not cured within five (5) business days after notice, or insolvency-related "
            "events. Upon an Event of Default, Provider may accelerate all unpaid amounts then due and suspend performance."
        ),
        "rationale": "Clear default mechanics make the transition from dispute to demand or filing much more supportable.",
        "jurisdiction_note": "Acceleration language can be scrutinized differently depending on contract type and state law.",
    },
    {
        "clause_type": "attorneys_fees",
        "title": "Attorneys' fees and collection costs",
        "patterns": [r"((?:attorneys'? fees|attorney fees|legal fees|collection costs|costs of collection)[\s\S]{0,260})"],
        "essential": True,
        "strong_terms": ["prevailing party", "all costs of collection", "enforcement"],
        "weak_terms": ["reasonable fees"],
        "dangerous_terms": [],
        "why": "Fee-shifting can materially change settlement leverage and whether recovery economics make sense.",
        "wrong": "If fees are missing or vague, legal spend may not be recoverable even if the claim is valid.",
        "improve": "Use prevailing-party or enforcement language that expressly covers collection costs and legal fees.",
        "fallback": (
            "The prevailing party in any action or proceeding to enforce this Agreement is entitled to recover its "
            "reasonable attorneys' fees, court costs, and collection expenses, whether incurred before suit, during litigation, or on appeal."
        ),
        "rationale": "Fee-shifting changes negotiation leverage and reduces the chance of an economically weak paper victory.",
        "jurisdiction_note": "Some states construe fee-shifting strictly, so wording should be checked for local enforceability.",
    },
    {
        "clause_type": "venue_forum",
        "title": "Venue and forum",
        "patterns": [r"((?:venue|forum|exclusive jurisdiction|courts of)[\s\S]{0,260})"],
        "essential": True,
        "strong_terms": ["exclusive", "county", "state", "jurisdiction"],
        "weak_terms": ["may be brought"],
        "dangerous_terms": [],
        "why": "Venue control helps prevent filing in a bad forum and reduces early procedural fights.",
        "wrong": "A vague venue clause can leave the dispute in a costly or inconvenient forum.",
        "improve": "Use an exclusive venue clause naming the county and state if that fits the business deal.",
        "fallback": (
            "Any action arising out of or relating to this Agreement must be brought exclusively in the state or federal "
            "courts located in the agreed county and state, and each party consents to personal jurisdiction and venue there."
        ),
        "rationale": "Forum clarity lowers filing risk and supports faster enforcement planning.",
        "jurisdiction_note": "Forum-selection enforceability can vary where consumer, employment, or public-policy limits apply.",
    },
    {
        "clause_type": "governing_law",
        "title": "Governing law",
        "patterns": [r"((?:governed by|governing law)[\s\S]{0,220})"],
        "essential": True,
        "strong_terms": ["laws of", "without regard to conflict"],
        "weak_terms": [],
        "dangerous_terms": [],
        "why": "Governing-law clauses reduce uncertainty about which rules apply to the contract.",
        "wrong": "If governing law is missing, parties can spend time and money arguing basic legal standards.",
        "improve": "Identify the governing state law and consider excluding conflict-of-laws rules.",
        "fallback": (
            "This Agreement is governed by and construed under the laws of the selected state, without regard to its conflict-of-laws rules."
        ),
        "rationale": "Predictable governing law improves drafting, collections strategy, and settlement valuation.",
        "jurisdiction_note": "Choice-of-law clauses can be limited where the chosen state lacks a sufficient relationship to the deal.",
    },
    {
        "clause_type": "notice",
        "title": "Notice mechanics",
        "patterns": [r"((?:notice|written notice|email notice|address for notice)[\s\S]{0,260})"],
        "essential": True,
        "strong_terms": ["email", "courier", "deemed received", "address for notice"],
        "weak_terms": ["written notice"],
        "dangerous_terms": [],
        "why": "Notice rules matter for default, termination, and evidence that the other side was warned.",
        "wrong": "If notice mechanics are unclear, the other side can dispute whether default or termination notices were effective.",
        "improve": "Specify valid notice methods, addresses, and when notice is deemed received.",
        "fallback": (
            "All notices must be in writing and sent by email with confirmation, nationally recognized courier, or certified mail "
            "to the notice addresses stated in this Agreement. Notice is deemed received on the date of confirmed delivery."
        ),
        "rationale": "Concrete notice mechanics make enforcement evidence much stronger.",
        "jurisdiction_note": "Some notice statutes or claim types can require special service methods beyond contract notice rules.",
    },
    {
        "clause_type": "termination",
        "title": "Termination rights",
        "patterns": [r"((?:termination|terminate)[\s\S]{0,260})"],
        "essential": True,
        "strong_terms": ["for cause", "for convenience", "survive", "payment due"],
        "weak_terms": ["terminate upon notice"],
        "dangerous_terms": ["without liability"],
        "why": "Termination language decides what survives if the relationship breaks down.",
        "wrong": "Bad termination language can leave earned fees or post-termination obligations unclear.",
        "improve": "Separate termination for cause and convenience, and state what payment or confidentiality obligations survive.",
        "fallback": (
            "Either party may terminate for material breach if the breach is not cured within ten (10) days after notice. "
            "Client remains responsible for all amounts accrued through the effective termination date, and payment, confidentiality, "
            "indemnity, and dispute-resolution provisions survive termination."
        ),
        "rationale": "Survival and accrued-payment language are common weak points in later collection disputes.",
        "jurisdiction_note": "Industry-specific cancellation rights may override contract language in some states.",
    },
    {
        "clause_type": "ach",
        "title": "ACH or payment authorization",
        "patterns": [r"((?:ACH|electronic withdrawal|debit authorization|autopay)[\s\S]{0,260})"],
        "essential": False,
        "strong_terms": ["authorize", "retry", "revoked in writing"],
        "weak_terms": [],
        "dangerous_terms": [],
        "why": "Payment authorization can reduce collection friction where recurring or scheduled payments are expected.",
        "wrong": "Without clear authorization terms, disputes over failed or repeated payments become harder to manage.",
        "improve": "Clarify authorization scope, retries, revocation mechanics, and returned-payment fees.",
        "fallback": (
            "Client authorizes Provider to initiate ACH debits for invoices due under this Agreement and to retry returned items once. "
            "Any revocation must be in writing and does not excuse previously accrued amounts."
        ),
        "rationale": "ACH language is often commercially useful for payment-plan and recurring-service contracts.",
        "jurisdiction_note": "Electronic payment authorization can implicate banking and consumer-debit rules depending on context.",
    },
    {
        "clause_type": "guaranty",
        "title": "Personal guaranty",
        "patterns": [r"((?:personal guaranty|personally guarantee|guarantor)[\s\S]{0,260})"],
        "essential": False,
        "strong_terms": ["personally guarantees", "joint and several", "individual capacity"],
        "weak_terms": ["guarantor"],
        "dangerous_terms": [],
        "why": "A real personal guaranty can materially improve collectability if the business counterparty is weak.",
        "wrong": "If signature capacity is unclear, the guaranty may fail because the signer can claim they signed only for the company.",
        "improve": "Make sure the guaranty is separate or clearly signed in an individual capacity.",
        "fallback": (
            "For value received, the undersigned individual unconditionally and personally guarantees payment and performance of all obligations "
            "under this Agreement, as a primary obligor and not merely as a surety."
        ),
        "rationale": "Collectability often depends more on a usable guaranty than on the underlying paper claim.",
        "jurisdiction_note": "Guaranty enforceability can turn on signature placement, separate consideration, and anti-deficiency limits.",
    },
    {
        "clause_type": "assignment",
        "title": "Assignment and transfer",
        "patterns": [r"((?:assignment|assign)[\s\S]{0,220})"],
        "essential": False,
        "strong_terms": ["without prior written consent", "successors and assigns"],
        "weak_terms": ["may assign"],
        "dangerous_terms": [],
        "why": "Assignment controls whether the other side can transfer the deal to an unknown party.",
        "wrong": "Loose assignment language can create payment and performance risk if obligations move to a new entity.",
        "improve": "Restrict assignment without consent, except for controlled internal restructurings if desired.",
        "fallback": (
            "Neither party may assign this Agreement without the other party's prior written consent, except to an affiliate or successor "
            "in a merger or sale of substantially all assets that assumes all obligations in writing."
        ),
        "rationale": "Assignment control protects operational risk and preserves contracting leverage.",
        "jurisdiction_note": "Some rights are assignable by law despite contract limits, especially payment rights.",
    },
    {
        "clause_type": "limitation_of_liability",
        "title": "Limitation of liability",
        "patterns": [r"((?:limitation of liability|liable for|consequential damages|exclusive remedy)[\s\S]{0,260})"],
        "essential": False,
        "strong_terms": [],
        "weak_terms": ["consequential damages"],
        "dangerous_terms": ["sole and exclusive remedy", "aggregate liability", "cap on liability"],
        "why": "Liability caps and exclusive-remedy language can wipe out claims that otherwise look strong.",
        "wrong": "A broad cap or exclusive remedy can leave one side with almost no practical recovery.",
        "improve": "Carve out payment obligations, fraud, confidentiality, indemnity, and intentional misconduct from any liability cap.",
        "fallback": (
            "Any limitation of liability does not apply to payment obligations, indemnity obligations, fraud, willful misconduct, "
            "confidentiality breaches, or amounts owed for delivered goods or services."
        ),
        "rationale": "This is a high-impact clause because it can override the economic value of the rest of the contract.",
        "jurisdiction_note": "Public-policy limits on liability waivers vary significantly by state and subject matter.",
    },
    {
        "clause_type": "indemnity",
        "title": "Indemnity",
        "patterns": [r"((?:indemnif(?:y|ication)|hold harmless|defend and indemnify)[\s\S]{0,260})"],
        "essential": False,
        "strong_terms": ["defend", "indemnify", "hold harmless"],
        "weak_terms": ["indemnify"],
        "dangerous_terms": [],
        "why": "Indemnity language allocates third-party risk and can materially change exposure.",
        "wrong": "One-sided or vague indemnity terms can create open-ended exposure or fail when needed most.",
        "improve": "Define who indemnifies whom, for what claims, and include procedure and carve-outs.",
        "fallback": (
            "Each party shall indemnify, defend, and hold the other harmless from third-party claims arising from its own negligence, "
            "willful misconduct, or breach of this Agreement, subject to prompt notice and reasonable cooperation."
        ),
        "rationale": "Indemnity is valuable only when the trigger, scope, and defense procedure are clear.",
        "jurisdiction_note": "Anti-indemnity statutes can limit these clauses in construction and other regulated industries.",
    },
    {
        "clause_type": "arbitration",
        "title": "Arbitration",
        "patterns": [r"((?:arbitration|arbitrate|binding arbitration)[\s\S]{0,260})"],
        "essential": False,
        "strong_terms": ["binding arbitration", "AAA", "rules"],
        "weak_terms": ["may arbitrate"],
        "dangerous_terms": ["exclusive arbitration"],
        "why": "Arbitration clauses can override litigation strategy and venue assumptions.",
        "wrong": "If arbitration is mandatory, filing in court can waste time and fees before the matter is even heard.",
        "improve": "Decide whether arbitration is actually desired and whether collections carve-outs are needed.",
        "fallback": (
            "Any arbitration clause should state the administering rules, seat, cost allocation, and whether claims for injunctive relief "
            "or collection of unpaid invoices may still be filed in court."
        ),
        "rationale": "Arbitration changes forum, cost, and timing, so it must be surfaced clearly.",
        "jurisdiction_note": "Arbitration enforceability is highly fact- and forum-sensitive, especially with statutory claims.",
    },
]


def _extract_clause(snippets: list[tuple[str, str]], patterns: list[str]) -> tuple[str | None, str | None]:
    for source_file, text in snippets:
        for pattern in patterns:
            match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if match:
                return _clean(match.group(1)), source_file
    return None, None


def _found_classification(spec: dict[str, Any], clause_text: str) -> tuple[str, str]:
    lower = str(clause_text or "").lower()
    if any(token in lower for token in spec.get("dangerous_terms", [])):
        return "dangerous", "high"
    if any(token in lower for token in spec.get("strong_terms", [])):
        return "strong", "low"
    if any(token in lower for token in spec.get("weak_terms", [])):
        return "weak", "medium"
    return "weak", "medium"


def _finding_from_spec(
    spec: dict[str, Any],
    clause_text: str | None,
    source_file: str | None,
) -> ContractClauseFinding:
    if clause_text:
        finding_kind, severity = _found_classification(spec, clause_text)
        practical_implication = spec["why"] if finding_kind == "strong" else spec["wrong"]
        recommended = "Keep this structure but confirm the exact wording fits the deal." if finding_kind == "strong" else spec["improve"]
        return ContractClauseFinding(
            clause_type=spec["clause_type"],
            title=spec["title"],
            finding_kind=finding_kind,
            extracted_text=clause_text,
            source_file=source_file,
            page_hint="review uploaded contract text",
            confidence=0.84,
            extraction_certainty=0.84,
            legal_certainty=0.46,
            risk_severity=severity,
            practical_implication=practical_implication,
            recommended_correction=recommended,
            why_it_matters=spec["why"],
            what_can_go_wrong=spec["wrong"],
            how_to_improve=spec["improve"],
            fallback_language=spec["fallback"],
            legal_rationale=spec["rationale"],
            jurisdiction_note=spec["jurisdiction_note"],
        )
    return ContractClauseFinding(
        clause_type=spec["clause_type"],
        title=spec["title"],
        finding_kind="missing",
        extracted_text=None,
        source_file=None,
        page_hint=None,
        confidence=0.0,
        extraction_certainty=0.0,
        legal_certainty=0.35,
        risk_severity="high" if spec.get("essential") else "medium",
        practical_implication=spec["wrong"],
        recommended_correction=spec["improve"],
        why_it_matters=spec["why"],
        what_can_go_wrong=spec["wrong"],
        how_to_improve=spec["improve"],
        fallback_language=spec["fallback"],
        legal_rationale=spec["rationale"],
        jurisdiction_note=spec["jurisdiction_note"],
    )


def analyze_contracts(record: CaseRecord) -> ContractReviewResult:
    findings: list[ContractClauseFinding] = []
    warnings: list[str] = []
    summary: list[str] = []
    jurisdiction_warnings: list[str] = []

    snippets = _combined_contract_text(record)
    if not snippets:
        return ContractReviewResult(status="blocked", warnings=["No contract or guaranty source text is available for review."])

    for spec in CLAUSE_LIBRARY:
        clause_text, source_file = _extract_clause(snippets, spec["patterns"])
        findings.append(_finding_from_spec(spec, clause_text, source_file))

    counts = {
        "strong": sum(1 for item in findings if item.finding_kind == "strong"),
        "weak": sum(1 for item in findings if item.finding_kind == "weak"),
        "missing": sum(1 for item in findings if item.finding_kind == "missing"),
        "dangerous": sum(1 for item in findings if item.finding_kind == "dangerous"),
    }

    summary.append(
        f"Review found {counts['strong']} strong, {counts['weak']} weak, {counts['missing']} missing, and {counts['dangerous']} dangerous clause positions."
    )
    if counts["missing"]:
        warnings.append("One or more important protections appear to be missing from the uploaded contract text.")
    if counts["dangerous"]:
        warnings.append("At least one clause appears to limit recovery or shift risk in a way that deserves careful review before signing.")
    if not counts["strong"]:
        warnings.append("Very few clauses look negotiation-ready in their current form.")

    for finding in findings:
        if finding.jurisdiction_note:
            jurisdiction_warnings.append(f"{finding.title}: {finding.jurisdiction_note}")

    return ContractReviewResult(
        status="completed" if findings else "review_required",
        findings=findings,
        summary=summary,
        warnings=warnings,
        jurisdiction_warnings=list(dict.fromkeys(jurisdiction_warnings)),
    )


def generate_revised_contract(record: CaseRecord, review_result: ContractReviewResult) -> ContractReviewResult:
    dirs = ensure_case_dirs(record.id)
    output_dir = dirs["output_dir"]

    revised_path = output_dir / "contract_review_revised.docx"
    memo_path = output_dir / "contract_review_change_memo.docx"

    doc = Document()
    doc.add_heading("Stronger Contract Draft", 0)
    doc.add_paragraph(
        "This draft preserves the apparent business relationship but replaces weak or missing legal mechanics with stronger fallback language. "
        "Human review is still required before signature."
    )
    doc.add_paragraph(f"Base matter: {record.title or record.reference or record.id}")
    doc.add_heading("Drafting priorities", level=1)
    for item in review_result.summary:
        doc.add_paragraph(item, style=None)

    for finding in review_result.findings:
        doc.add_heading(finding.title, level=2)
        doc.add_paragraph(f"Finding type: {finding.finding_kind} | Risk: {finding.risk_severity}")
        if finding.extracted_text:
            doc.add_paragraph(f"Current wording: {finding.extracted_text}")
        if finding.fallback_language:
            doc.add_paragraph(f"Stronger clause language: {finding.fallback_language}")
        if finding.how_to_improve:
            doc.add_paragraph(f"Drafting note: {finding.how_to_improve}")
        if finding.jurisdiction_note:
            doc.add_paragraph(f"State-specific caution: {finding.jurisdiction_note}")
    doc.save(revised_path)

    memo = Document()
    memo.add_heading("Detailed Change Memo", 0)
    memo.add_paragraph(
        "This memo explains why each clause was treated as strong, weak, missing, or dangerous, and what commercial or legal issue the stronger draft is trying to fix."
    )
    for finding in review_result.findings:
        memo.add_heading(finding.title, level=1)
        memo.add_paragraph(f"Finding type: {finding.finding_kind}")
        memo.add_paragraph(f"Risk severity: {finding.risk_severity}")
        if finding.why_it_matters:
            memo.add_paragraph(f"Why it matters: {finding.why_it_matters}")
        if finding.what_can_go_wrong:
            memo.add_paragraph(f"What can go wrong: {finding.what_can_go_wrong}")
        if finding.how_to_improve:
            memo.add_paragraph(f"How to improve it: {finding.how_to_improve}")
        if finding.legal_rationale:
            memo.add_paragraph(f"Legal / commercial rationale: {finding.legal_rationale}")
        if finding.fallback_language:
            memo.add_paragraph(f"Fallback language: {finding.fallback_language}")
        if finding.jurisdiction_note:
            memo.add_paragraph(f"Jurisdiction caution: {finding.jurisdiction_note}")
        if finding.extracted_text:
            memo.add_paragraph(f"Source excerpt: {finding.extracted_text}")
    memo.save(memo_path)

    review_result.revised_contract_path = str(revised_path)
    review_result.revised_contract_download_url = f"/files/cases/{record.id}/output/{revised_path.name}"
    review_result.change_memo_path = str(memo_path)
    review_result.change_memo_download_url = f"/files/cases/{record.id}/output/{memo_path.name}"
    review_result.generated_at = _utcnow()
    return review_result
