
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from threading import Lock, Thread
from uuid import uuid4

from docx import Document

from backend.models.case_schema import (
    CaseRecord,
    CaseStatus,
    GeneratedDocument,
    GenerationProgress,
    GenerationProgressItem,
    PacketGenerationResult,
)
from backend.services.audit_service import record_case_event
from backend.services.case_state import synchronize_case_status
from backend.services.case_store import ensure_case_dirs, load_case, save_case
from backend.services.docx_generator import generate_packet_documents
from backend.services.download_ledger import sync_generated_outputs
from backend.services.entitlements_service import record_usage_event
from backend.services.export_service import build_case_export
from backend.services.jurisdiction_engine import SUPPORTED_STATE_PACKS, analyze_case_jurisdiction
from backend.services.recovery_workspace_service import archive_packet_generation, append_communication_log, refresh_workspace_modules
from backend.services.validation_service import collect_case_warnings, collect_generation_blockers, template_coverage_report
from backend.services.template_registry import strict_templates

_active_jobs: dict[str, Thread] = {}
_job_lock = Lock()

ANALYSIS_ONLY_DOCUMENT_TYPES = {
    "jurisdiction_memo": "Jurisdiction Memo",
    "service_memo": "Service Path Memo",
    "evidence_checklist": "Evidence Checklist",
    "blocker_notice": "Missing Template Blocker Notice",
}
SAFE_NONPACK_DOCUMENT_TYPES = set(ANALYSIS_ONLY_DOCUMENT_TYPES.keys()) | {"demand_letter"}


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _progress_payload(selected_document_types: list[str], job_id: str) -> GenerationProgress:
    return GenerationProgress(
        job_id=job_id,
        status="queued",
        selected_document_types=list(selected_document_types),
        total_documents=len(selected_document_types),
        completed_documents=0,
        progress_percent=0.0,
        items=[GenerationProgressItem(document_type=item, document_name=item.replace("_", " ").title(), status="queued") for item in selected_document_types],
        started_at=_utcnow(),
    )


def _sync_progress_item(progress: GenerationProgress, document_type: str, document_name: str, status: str, generated_doc: GeneratedDocument | None = None, blocking_reason: str | None = None) -> None:
    existing = next((item for item in progress.items if item.document_type == document_type), None)
    if existing is None:
        existing = GenerationProgressItem(document_type=document_type, document_name=document_name)
        progress.items.append(existing)
    existing.document_name = document_name
    existing.status = status
    existing.blocking_reason = blocking_reason or (generated_doc.blocking_reason if generated_doc else None)
    existing.download_url = generated_doc.download_url if generated_doc else None


def _persist_progress(record: CaseRecord, progress: GenerationProgress) -> None:
    progress.total_documents = max(progress.total_documents, len(progress.selected_document_types))
    if progress.total_documents:
        progress.progress_percent = round((progress.completed_documents / progress.total_documents) * 100, 2)
    else:
        progress.progress_percent = 100.0
    record.generation_progress = progress
    save_case(record)


def _write_analysis_doc(case_id: str, output_dir: Path, document_type: str, title: str, body_lines: list[str]) -> GeneratedDocument:
    file_name = f"{document_type}.docx"
    path = output_dir / file_name
    doc = Document()
    doc.add_heading(title, 0)
    for line in body_lines:
        doc.add_paragraph(line)
    doc.save(path)
    return GeneratedDocument(
        document_name=title,
        document_path=str(path),
        document_type=document_type,
        download_url=f"/files/cases/{case_id}/output/{path.name}",
        status="generated",
    )


def _analysis_document_payloads(record: CaseRecord) -> dict[str, tuple[str, list[str]]]:
    analysis = record.jurisdiction_analysis or analyze_case_jurisdiction(record)
    record.jurisdiction_analysis = analysis
    recommended = analysis.recommended_forum
    defendant = record.extracted_case_data.defendant_entity.party_name or "Review required"
    amount = str(record.extracted_case_data.amount_reconciliation.case_summary_total or record.extracted_case_data.amount_in_controversy)
    common = [
        f"Matter: {record.title or record.reference or record.id}",
        f"Defendant: {defendant}",
        f"Supported amount: {amount}",
        "Human review required before court-facing release.",
    ]
    return {
        "jurisdiction_memo": (
            "Jurisdiction Memo",
            common + [
                f"Recommended forum: {recommended.label if recommended and recommended.label else 'Undetermined'}",
                *(analysis.reason_summary or ["No reliable forum ranking was available."]),
                *(["Blockers:"] + analysis.blockers if analysis.blockers else []),
            ],
        ),
        "service_memo": (
            "Service Path Memo",
            common + [
                *(analysis.service_notes or ["Service path remains review required."]),
                "Do not allege registered-agent service facts unless official verification supports them.",
            ],
        ),
        "evidence_checklist": (
            "Evidence Checklist",
            common + [
                "Required evidence to preserve:",
                "- Signed contract / guaranty",
                "- Invoice PDFs and totals",
                "- Reconciliation summary",
                "- Official registry captures where available",
            ],
        ),
        "blocker_notice": (
            "Missing Template Blocker Notice",
            common + [
                "No implemented court-pack is available for the recommended non-Florida forum in this build.",
                "The system generated only analysis-first outputs and blocked unsupported court-facing pleadings.",
            ],
        ),
    }


def _generate_analysis_documents(record: CaseRecord, output_dir: Path, selected_document_types: list[str], progress_callback=None) -> list[GeneratedDocument]:
    payloads = _analysis_document_payloads(record)
    docs: list[GeneratedDocument] = []
    total = len(selected_document_types)
    for index, doc_type in enumerate(selected_document_types, start=1):
        title, body_lines = payloads[doc_type]
        generated = _write_analysis_doc(record.id, output_dir, doc_type, title, body_lines)
        docs.append(generated)
        if progress_callback:
            progress_callback(doc_type, title, "generated", index, total, generated)
    return docs


def _packet_summary(
    selected_document_types: list[str],
    successful_docs: list[GeneratedDocument],
    blocked_docs: list[GeneratedDocument],
    not_applicable_docs: list[GeneratedDocument],
) -> str:
    requested_total = len(selected_document_types)
    parts: list[str] = []
    if successful_docs:
        parts.append(f"{len(successful_docs)} generated")
    if blocked_docs:
        parts.append(f"{len(blocked_docs)} blocked")
    if not_applicable_docs:
        parts.append(f"{len(not_applicable_docs)} not applicable")
    if not parts:
        parts.append("0 generated")
    return ", ".join(parts) + f" / {requested_total} requested"


def _non_pack_blockers(record: CaseRecord, selected_document_types: list[str]) -> list[str]:
    analysis = record.jurisdiction_analysis or analyze_case_jurisdiction(record)
    record.jurisdiction_analysis = analysis
    recommended = analysis.recommended_forum
    if not recommended or not recommended.forum_state:
        return []
    if recommended.forum_state in SUPPORTED_STATE_PACKS and recommended.state_pack_available:
        return []
    unsupported = [doc for doc in selected_document_types if doc not in SAFE_NONPACK_DOCUMENT_TYPES]
    if unsupported:
        return [
            f"Selected court-facing documents require a live {recommended.forum_state} state pack, but only analysis-first outputs are available in this build.",
            f"Blocked document types: {', '.join(unsupported)}",
        ]
    return []


def _run_job(case_id: str, selected_document_types: list[str], job_id: str) -> None:
    try:
        record = load_case(case_id)
        progress = _progress_payload(selected_document_types, job_id)
        progress.status = "generating"
        progress.message = "Document generation started."
        record.status = CaseStatus.GENERATING.value
        _persist_progress(record, progress)

        blockers = collect_generation_blockers(record, include_verification_gate=False)
        blockers.extend(_non_pack_blockers(record, selected_document_types))
        blockers = list(dict.fromkeys(blockers))
        if blockers:
            progress.status = "blocked"
            progress.message = "Packet generation blocked by missing required facts or unavailable state pack."
            progress.blocking_items = blockers
            progress.finished_at = _utcnow()
            _persist_progress(record, progress)
            record = load_case(case_id)
            synchronize_case_status(record)
            save_case(record)
            record_case_event(
                case_id,
                "generation.blocker",
                {"blocking_items": blockers, "status": str(record.status)},
            )
            return

        dirs = ensure_case_dirs(case_id)
        output_dir = dirs["output_dir"]

        def on_progress(document_type: str, document_name: str, status: str, completed: int, total: int, generated_doc: GeneratedDocument | None):
            live_record = load_case(case_id)
            live_progress = live_record.generation_progress or _progress_payload(selected_document_types, job_id)
            live_progress.status = "generating"
            live_progress.current_document_type = document_type
            live_progress.current_document_name = document_name
            live_progress.completed_documents = completed
            live_progress.total_documents = total
            live_progress.message = f"{document_name} — {status.replace('_', ' ')}"
            _sync_progress_item(live_progress, document_type, document_name, status, generated_doc)
            live_record.status = CaseStatus.GENERATING.value
            _persist_progress(live_record, live_progress)

        template_types = [item.document_type for item in strict_templates()]
        selected_template_docs = [item for item in selected_document_types if item in template_types]
        selected_analysis_docs = [item for item in selected_document_types if item in ANALYSIS_ONLY_DOCUMENT_TYPES]

        documents: list[GeneratedDocument] = []
        if selected_template_docs:
            documents.extend(
                generate_packet_documents(
                    record,
                    output_dir,
                    selected_document_types=selected_template_docs,
                    progress_callback=lambda doc_type, doc_name, status, completed, total, generated: on_progress(
                        doc_type, doc_name, status, completed, len(selected_document_types), generated
                    ),
                )
            )

        if selected_analysis_docs:
            already_completed = len(documents)
            analysis_docs = _generate_analysis_documents(record, output_dir, selected_analysis_docs)
            for offset, generated in enumerate(analysis_docs, start=1):
                on_progress(
                    generated.document_type,
                    generated.document_name,
                    "generated",
                    already_completed + offset,
                    len(selected_document_types),
                    generated,
                )
            documents.extend(analysis_docs)

        coverage = template_coverage_report(record=record, generated_docs=documents)
        successful_docs = [doc for doc in documents if str(doc.status).lower() == "generated"]
        not_applicable_docs = [doc for doc in documents if str(doc.status).lower() == "not_applicable"]
        blocked_docs = [doc for doc in documents if str(doc.status).lower() in {"blocked", "template_missing"}]
        failed_docs = [doc for doc in documents if str(doc.status).lower() not in {"generated", "not_applicable", "blocked", "template_missing"}]
        export_result = (
            build_case_export(
                record=record,
                output_dir=output_dir,
                generated_docs=documents,
                extra_warnings=collect_case_warnings(record),
            )
            if successful_docs
            else None
        )

        record = load_case(case_id)
        record.packet_generation = PacketGenerationResult(
            generated_at=export_result.generated_at if export_result else _utcnow(),
            documents=documents,
            zip_path=str(export_result.zip_path) if export_result else None,
            download_url=f"/files/cases/{case_id}/output/{export_result.zip_path.name}" if export_result else None,
            warnings=export_result.warnings if export_result else collect_case_warnings(record),
            selected_document_types=selected_document_types,
            summary=_packet_summary(selected_document_types, successful_docs, blocked_docs, not_applicable_docs),
        )
        if failed_docs:
            record.packet_generation.generation_qa_status = "failed"
        elif blocked_docs or record.packet_generation.warnings:
            record.packet_generation.generation_qa_status = "completed_with_warnings"
        else:
            record.packet_generation.generation_qa_status = "ready_for_review"
        record.packet_generation.release_state = "ready_for_review"
        record.packet_generation.review_required = True
        progress = record.generation_progress or _progress_payload(selected_document_types, job_id)
        if failed_docs:
            progress.status = "failed"
        elif blocked_docs and successful_docs:
            progress.status = "completed_with_blockers"
        elif blocked_docs and not successful_docs and not not_applicable_docs:
            progress.status = "blocked"
        else:
            progress.status = "completed"
        progress.message = record.packet_generation.summary
        progress.completed_documents = len(selected_document_types)
        progress.total_documents = len(selected_document_types)
        progress.current_document_name = None
        progress.current_document_type = None
        progress.finished_at = _utcnow()
        progress.blocking_items = list(dict.fromkeys([doc.blocking_reason for doc in blocked_docs if doc.blocking_reason]))
        for doc in documents:
            _sync_progress_item(progress, doc.document_type, doc.document_name, doc.status, doc, doc.blocking_reason)
        record.generation_progress = progress
        archive_entry = archive_packet_generation(record)
        if archive_entry:
            append_communication_log(
                record,
                entry_type="generated_draft",
                channel="packet_archive",
                direction="outbound_draft",
                summary=f"Generated packet v{archive_entry.version_number} for review.",
                body=record.packet_generation.summary,
                actor_label="Generation engine",
                related_packet_version=archive_entry.version_number,
                metadata={
                    "release_state": archive_entry.release_state,
                    "document_types": archive_entry.document_types,
                    "selected_document_types": archive_entry.selected_document_types,
                    "is_court_draft": archive_entry.is_court_draft,
                },
            )
        refresh_workspace_modules(record)
        sync_generated_outputs(record)

        if failed_docs:
            record.status = CaseStatus.GENERATION_FAILED.value
            record_case_event(
                case_id,
                "document.qa.failed",
                {
                    "failed_document_count": len(failed_docs),
                    "failed_documents": [doc.model_dump(mode="json") for doc in failed_docs],
                    "selected_document_types": selected_document_types,
                },
            )
        elif blocked_docs and not successful_docs and not not_applicable_docs:
            record.status = CaseStatus.BLOCKED.value
            record_case_event(
                case_id,
                "generation.blocker",
                {
                    "blocked_document_count": len(blocked_docs),
                    "blocked_documents": [doc.model_dump(mode="json") for doc in blocked_docs],
                    "selected_document_types": selected_document_types,
                },
            )
        else:
            record.status = CaseStatus.PACKET_GENERATED.value
            record_case_event(
                case_id,
                "generation.partial" if blocked_docs else "generation.success",
                {
                    "document_count": len(successful_docs),
                    "blocked_document_count": len(blocked_docs),
                    "not_applicable_document_count": len(not_applicable_docs),
                    "zip_download_url": record.packet_generation.download_url,
                    "warnings": record.packet_generation.warnings,
                    "coverage": coverage,
                    "selected_document_types": selected_document_types,
                },
            )
            record_usage_event(
                record.workspace_id or "workspace_legacy",
                "packet_generated",
                case_id=case_id,
                actor_user_id=record.created_by_user_id,
                payload={
                    "selected_document_types": selected_document_types,
                    "blocked_document_count": len(blocked_docs),
                    "successful_document_count": len(successful_docs),
                },
            )
            if "demand_letter" in selected_document_types:
                record_usage_event(
                    record.workspace_id or "workspace_legacy",
                    "demand_pack_created",
                    case_id=case_id,
                    actor_user_id=record.created_by_user_id,
                    payload={"selected_document_types": selected_document_types},
                )

        synchronize_case_status(record)
        save_case(record)
    except Exception as exc:
        try:
            record = load_case(case_id)
            progress = record.generation_progress or _progress_payload(selected_document_types, job_id)
            progress.status = "failed"
            progress.message = str(exc)
            progress.finished_at = _utcnow()
            record.generation_progress = progress
            record.status = CaseStatus.GENERATION_FAILED.value
            save_case(record)
            record_case_event(case_id, "packet.generation.failed", {"error": str(exc)})
        except Exception:
            pass
        raise
    finally:
        with _job_lock:
            _active_jobs.pop(case_id, None)


def start_generation_job(case_id: str, selected_document_types: list[str]) -> GenerationProgress:
    with _job_lock:
        existing = _active_jobs.get(case_id)
        if existing and existing.is_alive():
            record = load_case(case_id)
            return record.generation_progress or _progress_payload(selected_document_types, "running")
        job_id = uuid4().hex
        record = load_case(case_id)
        progress = _progress_payload(selected_document_types, job_id)
        record.generation_progress = progress
        record.packet_generation = None
        record.status = CaseStatus.GENERATING.value
        save_case(record)
        thread = Thread(target=_run_job, args=(case_id, selected_document_types, job_id), daemon=True)
        _active_jobs[case_id] = thread
        thread.start()
        return progress


def run_generation_job_sync(case_id: str, selected_document_types: list[str]) -> GenerationProgress:
    job_id = uuid4().hex
    record = load_case(case_id)
    record.generation_progress = _progress_payload(selected_document_types, job_id)
    record.packet_generation = None
    record.status = CaseStatus.GENERATING.value
    save_case(record)
    _run_job(case_id, selected_document_types, job_id)
    record = load_case(case_id)
    return record.generation_progress or _progress_payload(selected_document_types, job_id)


def get_generation_progress(case_id: str) -> GenerationProgress:
    record = load_case(case_id)
    return record.generation_progress or GenerationProgress(status="idle")
